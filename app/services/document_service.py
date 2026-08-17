from html import unescape
from html.parser import HTMLParser
from io import BytesIO
import math
import re
import struct
import zlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from flask import make_response


class HtmlToDocxParser(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.current_paragraph = None
        self.current_href = None
        self.bold_depth = 0
        self.list_stack: list[str] = []
        self.in_list_item = False

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag == "h1":
            self._start_paragraph(style="Heading 1")
        elif tag == "h2":
            self._start_paragraph(style="Heading 2")
        elif tag == "h3":
            self._start_paragraph(style="Heading 3")
        elif tag == "p":
            self._start_paragraph()
        elif tag == "ul":
            self.list_stack.append("bullet")
        elif tag == "ol":
            self.list_stack.append("number")
        elif tag == "li":
            style = "List Bullet" if self._current_list_type() == "bullet" else "List Number"
            self._start_paragraph(style=style)
            self.in_list_item = True
        elif tag == "br":
            if self.current_paragraph is None:
                self._start_paragraph()
            self.current_paragraph.add_run().add_break()
        elif tag in {"strong", "b"}:
            self.bold_depth += 1
        elif tag == "a":
            self.current_href = attrs_dict.get("href", "")

    def handle_endtag(self, tag):
        if tag in {"h1", "h2", "h3", "p", "li"}:
            self.current_paragraph = None
        if tag == "li":
            self.in_list_item = False
        elif tag in {"ul", "ol"} and self.list_stack:
            self.list_stack.pop()
        elif tag in {"strong", "b"} and self.bold_depth > 0:
            self.bold_depth -= 1
        elif tag == "a":
            self.current_href = None

    def handle_data(self, data):
        text = unescape(data)
        if not text:
            return

        if self.current_paragraph is None:
            if not text.strip():
                return
            self._start_paragraph()

        if self.current_href:
            self._add_hyperlink(self.current_paragraph, text, self.current_href, bold=self.bold_depth > 0)
            return

        run = self.current_paragraph.add_run(text)
        if self.bold_depth > 0:
            run.bold = True

    def _start_paragraph(self, style=None):
        self.current_paragraph = self.doc.add_paragraph(style=style)
        return self.current_paragraph

    def _current_list_type(self):
        if not self.list_stack:
            return None
        return self.list_stack[-1]

    def _add_hyperlink(self, paragraph, text: str, url: str, bold: bool = False):
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        part = paragraph.part
        relation_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relation_id)

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        run_style = OxmlElement("w:rStyle")
        run_style.set(qn("w:val"), "Hyperlink")
        run_properties.append(run_style)

        if bold:
            bold_element = OxmlElement("w:b")
            run_properties.append(bold_element)

        run.append(run_properties)

        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)


def html_to_docx_paragraph(doc, html_content):
    cleaned = (html_content or "").strip()
    if not cleaned:
        return

    parser = HtmlToDocxParser(doc)
    parser.feed(cleaned)
    parser.close()


def extract_docx_text(file_storage, max_chars: int = 12000) -> str:
    filename = (getattr(file_storage, "filename", "") or "").strip()
    if not filename.casefold().endswith(".docx"):
        raise ValueError("Please upload a DOCX reference file.")
    file_storage.stream.seek(0)
    doc = Document(BytesIO(file_storage.read()))
    parts = []
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()).strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts).strip()
    return text[:max_chars]


def extract_docx_website_reference(file_storage, max_chars: int = 12000) -> dict:
    filename = (getattr(file_storage, "filename", "") or "").strip()
    if not filename.casefold().endswith(".docx"):
        raise ValueError("Please upload a DOCX reference file.")
    file_storage.stream.seek(0)
    doc = Document(BytesIO(file_storage.read()))
    parts = []
    pages = []
    current_page = None

    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue
        parts.append(text)
        style_name = (paragraph.style.name if paragraph.style else "").casefold()
        label, labeled_value = _docx_labeled_value(text)

        if "heading 1" in style_name or label in {"page", "page name"}:
            current_page = _docx_page(labeled_value or text)
            if "heading 1" in style_name:
                current_page["h1"] = text
            pages.append(current_page)
            continue

        if label == "h1":
            if current_page is None:
                current_page = _docx_page(labeled_value)
                pages.append(current_page)
            current_page["h1"] = labeled_value
            continue

        if label == "keyword":
            if current_page is None:
                current_page = _docx_page(labeled_value)
                pages.append(current_page)
            current_page["keyword"] = labeled_value
            current_page["name"] = labeled_value
            continue

        if "heading " in style_name or label in {"h2", "h3", "heading"}:
            if current_page is None:
                current_page = _docx_page("Reference Page")
                pages.append(current_page)
            heading = labeled_value or text
            if heading and heading not in current_page["headings"]:
                current_page["headings"].append(heading)

    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()).strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    cleaned_pages = []
    seen = set()
    for page in pages:
        keyword = " ".join((page.get("keyword") or page.get("name") or page.get("h1") or "").split()).strip()
        h1 = " ".join((page.get("h1") or page.get("name") or keyword).split()).strip()
        if not keyword:
            continue
        key = keyword.casefold()
        if key in seen:
            continue
        seen.add(key)
        headings = _unique_text(page.get("headings", []))
        cleaned_pages.append({"name": keyword, "keyword": keyword, "h1": h1, "headings": headings})

    return {"text": "\n".join(parts).strip()[:max_chars], "pages": cleaned_pages}


def build_docx_response(
    title: str,
    keyword: str,
    supporting_keyword: str,
    meta_description: str,
    content_html: str,
    medium_name: str = "",
    tags: str = "",
):
    doc = Document()

    title_para = doc.add_paragraph(title, style="Heading 1")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    metadata_paragraph = doc.add_paragraph()
    metadata_paragraph.add_run("Keyword: ").bold = True
    metadata_paragraph.add_run(keyword)

    if supporting_keyword:
        supporting_paragraph = doc.add_paragraph()
        supporting_paragraph.add_run("Supporting Keyword: ").bold = True
        supporting_paragraph.add_run(supporting_keyword)

    if medium_name:
        medium_paragraph = doc.add_paragraph()
        medium_paragraph.add_run("Medium: ").bold = True
        medium_paragraph.add_run(medium_name)

    meta_paragraph = doc.add_paragraph()
    meta_paragraph.add_run("Meta Description: ").bold = True
    meta_paragraph.add_run(meta_description)

    if tags:
        tags_paragraph = doc.add_paragraph()
        tags_paragraph.add_run("Tags: ").bold = True
        tags_paragraph.add_run(tags)

    doc.add_paragraph()
    html_to_docx_paragraph(doc, content_html)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    filename = title.replace(" ", "_")[:50] or "blog_post"
    response = make_response(doc_io.getvalue())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
    return response


def build_gsc_summary_report_response(
    brand: str,
    target_url: str,
    gsc_property: str,
    start_date: str,
    end_date: str,
    report: dict,
    query_rows: list[dict] | None = None,
    daily_rows: list[dict] | None = None,
    backlink_snapshot: dict | None = None,
):
    doc = Document()
    title = f"GSC Summary Report - {(brand or 'SEO Report').strip()}"

    title_para = doc.add_paragraph(title, style="Heading 1")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_label_value(doc, "Brand", brand)
    _add_label_value(doc, "Target URL", target_url)
    _add_label_value(doc, "Search Console Property", gsc_property)
    _add_label_value(doc, "Date Range", " to ".join(item for item in (start_date, end_date) if item))

    summary = (report or {}).get("executive_summary", "")
    if summary:
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(summary)

    if daily_rows:
        doc.add_heading("Performance Trend Graph", level=2)
        trend_image = _gsc_trend_png(daily_rows)
        if trend_image:
            doc.add_picture(trend_image, width=Inches(6.6))
            doc.add_paragraph("Green line: clicks. Brown line: impressions. Each series is scaled independently so changes are visible.")
        _add_gsc_totals_table(doc, daily_rows)

    if query_rows:
        doc.add_heading("Top Query Graph", level=2)
        _add_query_bar_table(doc, query_rows[:10])

    backlink_snapshot = backlink_snapshot or {}
    if backlink_snapshot:
        doc.add_heading("Backlink Snapshot", level=2)
        doc.add_paragraph(
            f"{backlink_snapshot.get('total_count', 0)} saved backlink(s), "
            f"{backlink_snapshot.get('scored_count', 0)} with DP/DA/DR, "
            f"{backlink_snapshot.get('unscored_count', 0)} without score."
        )

    for title_text, rows, keys in _gsc_report_sections(report or {}):
        if not rows:
            continue
        doc.add_heading(title_text, level=2)
        _add_section_table(doc, rows, keys)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    filename = _download_filename(title)
    response = make_response(doc_io.getvalue())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
    return response


def build_website_planner_report_response(metadata: dict, plan: dict):
    doc = Document()
    client = (metadata or {}).get("client", "").strip()
    domain = (metadata or {}).get("domain", "").strip()
    title_subject = client or domain or "Website Planner"
    title_para = doc.add_paragraph(f"SEO Master Plan - {title_subject}", style="Heading 1")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    main_pages = _planner_items(plan, "main_pages")
    trust_pages = _planner_items(plan, "trust_pages")
    blogs = _planner_items(plan, "blogs")
    target_market = (metadata or {}).get("target_market", "").strip()
    language = (metadata or {}).get("language", "").strip()
    site_type = (metadata or {}).get("site_type", "").strip()
    reference_content = (metadata or {}).get("reference_content", "").strip()
    reference_filename = (metadata or {}).get("reference_filename", "").strip()

    doc.add_heading("About site", level=2)
    _add_simple_table(
        doc,
        ("Field", "Details"),
        [
            ("Client", client),
            ("Domain", domain),
            ("Target Market", target_market),
            ("Language", language),
            ("Site Type", site_type),
            ("Date", ""),
        ],
    )
    if reference_content:
        doc.add_heading("Website Content Reference", level=2)
        if reference_filename:
            _add_label_value(doc, "Reference File", reference_filename)
        doc.add_paragraph(_truncate_text(reference_content, 3000))

    doc.add_heading("PART 1 - KEYWORD PLAN", level=2)
    doc.add_heading("Strategic Overview", level=3)
    doc.add_paragraph(
        "Use this keyword plan to map one primary search intent to each core page, then support it with secondary terms and blog topics. "
        "Avoid claims such as official, licensed, or approved unless those claims are verified for the site."
    )
    _add_step_prompt(
        doc,
        [
            "Review the client, domain, market, language, and site type before choosing commercial or informational intent.",
            "Read the uploaded website content reference and reuse factual details, page themes, offers, and terminology where relevant.",
            "Assign one main keyword to every core page and keep trust pages out of the keyword clusters.",
            "Add volume, competition, difficulty, and intent from your keyword research tool before final publishing.",
            "Use the finalized keyword map as the source for page titles, H2s, FAQs, and internal anchors.",
        ],
    )

    doc.add_heading("Keyword CLUSTER", level=3)
    for index, page in enumerate(main_pages, start=1):
        keyword = _planner_page_keyword(page, f"Main Page {index}")
        page_name = _planner_page_name(page, keyword)
        h1 = _planner_page_h1(page, keyword)
        doc.add_paragraph(f"CLUSTER {index}: {page_name}", style="Heading 3")
        _add_label_value(doc, "Primary Target Page", _planner_page_slug(page, keyword))
        _add_label_value(doc, "Main Keyword", keyword)
        _add_label_value(doc, "H1 / Page Title", h1)
        if page.get("headings"):
            _add_label_value(doc, "Extracted Headings", ", ".join(page.get("headings", [])))
        _add_simple_table(
            doc,
            ("Keyword", "Monthly Volume", "Competition", "Keyword Difficulty", "Intent"),
            _planner_keyword_rows(keyword, client, target_market),
        )

    doc.add_heading("PART 2 - CONTENT LAYOUT PLAN", level=2)
    doc.add_heading("CONTENT LAYOUT PLAN", level=3)
    _add_step_prompt(
        doc,
        [
            "Start each core page with a clear search-intent match: problem, offer, benefit, and next action.",
            "Pull useful phrasing, service details, audience notes, and proof points from the uploaded reference content.",
            "Convert the page keyword cluster into H2 sections, FAQs, examples, and conversion blocks.",
            "Place support content below the main conversion path so visitors can scan the page quickly.",
            "Add one internal link from every major section to the most relevant core page or blog support topic.",
        ],
    )
    doc.add_heading("Site Architecture (Silo Structure)", level=3)
    doc.add_paragraph("Homepage (/)")
    for page in main_pages:
        keyword = _planner_page_keyword(page)
        doc.add_paragraph(f"{_planner_page_slug(page, keyword)} [{_planner_page_h1(page, keyword)}]", style="List Bullet")
    doc.add_paragraph("/blog/ [Blog Hub]", style="List Bullet")
    for blog in blogs[:10]:
        doc.add_paragraph(f"/blog/{_slug_text(blog.get('name', 'topic'))}/ [{blog.get('name', '')}]", style="List Bullet")

    for index, page in enumerate(main_pages, start=1):
        keyword = _planner_page_keyword(page, f"Main Page {index}")
        page_name = _planner_page_name(page, keyword)
        h1 = _planner_page_h1(page, keyword)
        headings = page.get("headings", []) if isinstance(page.get("headings", []), list) else []
        doc.add_heading(f"Page {index}: {page_name}", level=3)
        _add_label_value(doc, "Primary Keyword", keyword)
        _add_label_value(doc, "H1 / Page Title", h1)
        if headings:
            _add_label_value(doc, "Extracted Headings", ", ".join(headings))
        _add_label_value(doc, "Secondary Keywords", ", ".join(_planner_secondary_keywords(keyword, client, target_market)))
        _add_simple_table(
            doc,
            ("Section", "Purpose", "Prompt"),
            [
                ("Hero Section", "Clarify offer and primary intent", f"Write a concise hero for {h1} using {keyword}."),
                ("H2 Guide Section", "Answer the main search question", _planner_heading_prompt(keyword, headings, target_market)),
                ("Benefits / Features", "Support conversion", f"List practical benefits, proof points, and limitations for {h1}."),
                ("FAQ", "Capture long-tail questions", f"Write 4-6 FAQs based on {keyword} and related search intent."),
                ("CTA", "Move the user forward", "Add one clear next step that matches the page intent."),
            ],
        )

    doc.add_heading("PART 3 - INTERNAL LINKING PLAN", level=2)
    _add_step_prompt(
        doc,
        [
            "Use the homepage as the main hub that links to every core page.",
            "Link each blog topic back to the most relevant core page with natural anchor text.",
            "Add contextual links between related core pages where the user journey naturally continues.",
            "Review anchors monthly so exact-match phrases do not become repetitive.",
        ],
    )
    doc.add_heading("Internal Link Map", level=3)
    _add_simple_table(doc, ("Source Page", "Link To", "Suggested Anchor", "Purpose"), _planner_internal_link_rows(main_pages, blogs))

    doc.add_heading("Anchor Text Guidelines", level=3)
    for guideline in (
        "Use descriptive anchors that explain what the reader will get after clicking.",
        "Mix exact-match, partial-match, branded, and plain-language anchors.",
        "Avoid using the same exact anchor repeatedly across every page.",
        "Prioritize links that help users move from research pages to core conversion pages.",
    ):
        doc.add_paragraph(guideline, style="List Bullet")

    doc.add_heading("PART 4 - BLOG CONTENT CATEGORIES AND CURRENT PAGE-DERIVED TOPICS", level=2)
    _add_step_prompt(
        doc,
        [
            "Group blog ideas by the core page they support.",
            "Choose topics that answer questions the core page cannot cover in full.",
            "Add a contextual link from each blog to the mapped core page.",
            "Publish supporting topics before or shortly after the matching core page goes live.",
        ],
    )
    _add_simple_table(
        doc,
        ("Category / Topic", "Source", "Target Core Page", "Suggested Intent", "Prompt"),
        _planner_blog_rows(main_pages, blogs),
    )

    doc.add_heading("PART 5 - SUMMARY DASHBOARD", level=2)
    doc.add_heading("SUMMARY DASHBOARD", level=3)
    summary = plan.get("summary", {}) if isinstance(plan.get("summary", {}), dict) else {}
    _add_simple_table(
        doc,
        ("Metric", "Value"),
        [
            ("Main Pages", str(summary.get("main_pages", len(main_pages)))),
            ("Trust Pages", str(summary.get("trust_pages", len(trust_pages)))),
            ("Blog Topics", str(summary.get("blogs", len(blogs)))),
            ("Total Items", str(summary.get("total", len(main_pages) + len(trust_pages) + len(blogs)))),
            ("Keyword Source", str(summary.get("blog_keyword_source", ""))),
            ("Target Market", target_market),
            ("Language", language),
        ],
    )
    _add_step_prompt(
        doc,
        [
            "Use the dashboard to confirm scope before writing or assigning content.",
            "Check that every core page has a primary keyword, content role, and supporting blog topics.",
            "Move pages with revenue or conversion value into the earliest launch days.",
        ],
    )

    doc.add_heading("Current Core Page Mapping", level=3)
    _add_simple_table(
        doc,
        ("Core Page", "URL Slug", "Primary Keyword", "Content Role", "Linked Blog Topics"),
        _planner_core_mapping_rows(main_pages, blogs),
    )

    doc.add_heading("PART 6 - PRIORITY LAUNCH ORDER (First 7 Days)", level=2)
    _add_step_prompt(
        doc,
        [
            "Launch the homepage and highest-value core pages first.",
            "Publish trust pages early enough to support user confidence and compliance.",
            "Add supporting blogs after their target core page exists.",
            "Finish the first week by checking internal links, metadata, indexability, and analytics.",
        ],
    )
    _add_simple_table(doc, ("Day", "Priority Work", "Output", "Notes"), _planner_launch_rows(main_pages, trust_pages, blogs))

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    filename = _download_filename(f"Website Planner Report {title_subject}")
    response = make_response(doc_io.getvalue())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
    return response


def build_seo_checker_report_response(result: dict):
    result = result or {}
    mode = result.get("mode", "single")
    title_subject = result.get("base_url") or result.get("url") or result.get("source_url") or "Website SEO Checker"
    title = f"Website SEO Checker Report - {title_subject}"
    doc = Document()
    title_para = doc.add_paragraph(title, style="Heading 1")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if mode == "site":
        summary = result.get("summary", {}) or {}
        doc.add_heading("Summary Dashboard", level=2)
        _add_simple_table(
            doc,
            ("Metric", "Value"),
            [
                ("Base URL", result.get("base_url", "")),
                ("Average Grade", summary.get("average_grade", "")),
                ("Average Score", f"{summary.get('average_score', 0)}/100"),
                ("Pages Listed", summary.get("discovered_count", 0)),
                ("Pages Checked", summary.get("checked_count", 0)),
                ("Warning/Fix Items", summary.get("issue_count", 0)),
                ("Checked Links", summary.get("checked_link_count", 0)),
                ("404 Links", summary.get("not_found_link_count", 0)),
                ("Redirecting Links", summary.get("redirect_link_count", 0)),
            ],
        )
        _add_seo_page_scores(doc, result.get("pages", []))
        _add_seo_link_health(doc, result.get("link_health_details", {}) or {})
        _add_seo_page_issues(doc, result.get("pages", []))
    else:
        doc.add_heading("Summary Dashboard", level=2)
        stats = result.get("stats", {}) or {}
        links = stats.get("links", {}) or {}
        _add_simple_table(
            doc,
            ("Metric", "Value"),
            [
                ("URL", result.get("url", "")),
                ("Grade", result.get("grade", "")),
                ("Score", f"{result.get('score', 0)}/100"),
                ("Status Code", result.get("status_code", "")),
                ("Title", stats.get("title", "")),
                ("Meta Description", stats.get("meta_description", "")),
                ("Word Count", stats.get("word_count", 0)),
                ("Missing Alt Text", stats.get("missing_alt_count", 0)),
                ("Checked Links", links.get("checked_count", 0)),
                ("404 Links", links.get("not_found_count", 0)),
                ("Redirecting Links", links.get("redirect_count", 0)),
            ],
        )
        _add_seo_link_health(doc, links)
        _add_seo_checks(doc, result.get("checks", []))

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    filename = _download_filename(title)
    response = make_response(doc_io.getvalue())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
    return response


def build_website_planner_v2_detailed_report_response(metadata: dict, plan: dict):
    doc = Document()
    client = (metadata or {}).get("client", "").strip()
    domain = (metadata or {}).get("domain", "").strip()
    title_subject = client or domain or "Website Planner V2"
    title_para = doc.add_paragraph(f"SEO Master Plan - {title_subject}", style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    core_pages = _planner_items(plan, "core_pages")
    support_pages = _planner_items(plan, "support_pages")
    blogs = _planner_items(plan, "blogs")
    categories = _planner_items(plan, "categories")
    blog_category_plans = _planner_items(plan, "blog_category_plans")
    summary = plan.get("summary", {}) if isinstance(plan.get("summary", {}), dict) else {}
    target_market = (metadata or {}).get("target_market", "").strip()
    language = (metadata or {}).get("language", "").strip()
    site_type = (metadata or {}).get("site_type", "").strip()
    source_filename = (metadata or {}).get("reference_filename", "").strip()

    doc.add_paragraph(
        " ".join(
            item
            for item in (
                f"Client: {client}" if client else "",
                f"Domain: {domain}" if domain else "",
                f"Target Market: {target_market}" if target_market else "",
                f"Language: {language}" if language else "",
                f"Site Type: {site_type}" if site_type else "",
                "Date:",
            )
            if item
        )
    )
    if source_filename:
        _add_label_value(doc, "Keyword Source File", source_filename)

    doc.add_heading("PART 1 - KEYWORD PLAN", level=2)
    doc.add_paragraph(
        "Strategic Overview: build one focused core page for each major keyword category, then support it with clear internal links, "
        "blog topics, FAQ coverage, and conversion sections. Avoid using unverified trust claims such as official, licensed, legit, "
        "or approved unless the client has proof for those statements."
    )
    for index, page in enumerate(core_pages, start=1):
        keyword = page.get("primary_keyword", "") or _planner_page_keyword(page, f"Core Page {index}")
        page_name = page.get("name", f"Core Page {index}")
        doc.add_paragraph(f"CLUSTER {index}: {page_name}", style="Heading 3")
        _add_label_value(doc, "Primary Target Page", page.get("slug", _planner_slug(keyword)))
        _add_label_value(doc, "Main Keyword", keyword)
        _add_label_value(doc, "H1 / Page Title", page.get("h1", keyword))
        _add_simple_table(
            doc,
            ("Keyword", "Monthly Volume", "Global Volume", "Traffic Potential", "Keyword Difficulty", "Intents"),
            _v2_detailed_keyword_rows(page),
        )

    doc.add_heading("PART 2 - CONTENT LAYOUT PLAN", level=2)
    doc.add_paragraph("Site Architecture (Silo Structure)")
    doc.add_paragraph("Homepage (/)")
    for page in core_pages:
        slug = page.get("slug", _planner_slug(page.get("primary_keyword", page.get("name", "page"))))
        doc.add_paragraph(f"{slug} [{page.get('name', '')}]", style="List Bullet")
    doc.add_paragraph("/blog/ [Blog Hub]", style="List Bullet")
    for category in categories[:10]:
        doc.add_paragraph(f"/blog/category/{_slug_text(category.get('name', 'category'))}/", style="List Bullet")

    for index, page in enumerate(core_pages, start=1):
        _add_v2_detailed_page_layout(doc, index, page, target_market)

    doc.add_heading("PART 3 - INTERNAL LINKING PLAN", level=2)
    doc.add_paragraph("Use the homepage as the main hub. Link core pages to related categories and link every generated blog topic back to the most relevant target page.")
    doc.add_heading("Internal Link Map", level=3)
    _add_v2_detailed_internal_link_map(doc, core_pages, support_pages, blogs)
    doc.add_heading("Anchor Text Guidelines", level=3)
    for guideline in (
        "Use natural anchors that match the destination page topic.",
        "Mix exact-match, partial-match, branded, and CTA-style anchors.",
        "Avoid repeating the same anchor across every page.",
        "Use blog posts to support core commercial pages without forcing the link.",
    ):
        doc.add_paragraph(guideline, style="List Bullet")

    doc.add_heading("PART 4 - BLOG CONTENT CATEGORIES AND CURRENT PAGE-DERIVED TOPICS", level=2)
    _add_simple_table(
        doc,
        ("#", "Topic", "Target Keyword", "Intent", "Category", "Target Core Page", "Source"),
        _v2_detailed_blog_rows(blogs),
    )
    doc.add_heading("Blog Category Title And Meta Description Plan", level=3)
    _add_simple_table(
        doc,
        ("Category", "Category Title", "Title Characters", "Meta Description", "Meta Characters", "Topics"),
        _v2_detailed_blog_category_rows(blog_category_plans),
    )

    doc.add_heading("PART 5 - SUMMARY DASHBOARD", level=2)
    _add_simple_table(
        doc,
        ("Plan Section", "Number of Pages or Items", "Current Role"),
        [
            ("Core Pages", summary.get("core_page_count", len(core_pages)), "Main SEO landing pages from keyword categories"),
            ("Support Pages", summary.get("support_page_count", len(support_pages)), "Trust, compliance, and contact pages"),
            ("Blog Topics", summary.get("blog_count", len(blogs)), summary.get("blog_keyword_source", "Blog support plan")),
            ("Keyword Categories", summary.get("category_count", len(categories)), "Category groups used for core page planning"),
            ("Total Search Volume", summary.get("total_volume", ""), "Imported Ahrefs keyword volume"),
        ],
    )
    doc.add_heading("Current Core Page Mapping", level=3)
    _add_simple_table(
        doc,
        ("Core Page", "URL Slug", "Primary Keyword", "H1", "Linked Blog Topics"),
        _v2_detailed_core_mapping_rows(core_pages, blogs),
    )

    doc.add_heading("PART 6 - PRIORITY LAUNCH ORDER (First 7 Days)", level=2)
    _add_simple_table(doc, ("Day", "Priority Work", "Output", "Notes"), _v2_detailed_launch_rows(core_pages, support_pages, blogs))

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    filename = _download_filename(f"Website Planner V2 Detailed Report {title_subject}")
    response = make_response(doc_io.getvalue())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
    return response


def _add_label_value(doc: Document, label: str, value: str) -> None:
    if not value:
        return
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value))


def _docx_labeled_value(text: str) -> tuple[str, str]:
    match = re.match(r"^\s*(page\s+name|page|keyword|h1|h2|h3|heading)\s*[:\-]\s*(.+?)\s*$", text or "", flags=re.IGNORECASE)
    if not match:
        return "", ""
    return match.group(1).casefold(), match.group(2).strip()


def _docx_page(value: str) -> dict:
    cleaned = " ".join(str(value or "Reference Page").split()).strip() or "Reference Page"
    return {"name": cleaned, "keyword": cleaned, "h1": cleaned, "headings": []}


def _add_step_prompt(doc: Document, steps: list[str]) -> None:
    doc.add_paragraph("Step-by-step prompt:", style="Heading 3")
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def _add_seo_page_scores(doc: Document, pages: list[dict]) -> None:
    rows = []
    for page in pages or []:
        if page.get("status") == "checked" and page.get("result"):
            rows.append(
                (
                    page.get("url", ""),
                    page.get("result", {}).get("grade", ""),
                    f"{page.get('result', {}).get('score', 0)}/100",
                    page.get("result", {}).get("stats", {}).get("title", ""),
                )
            )
        else:
            rows.append((page.get("url", ""), "Error", "", page.get("error", "")))
    if rows:
        doc.add_heading("Page Scores", level=2)
        _add_simple_table(doc, ("Page", "Grade", "Score", "Title / Error"), rows)


def _add_seo_page_issues(doc: Document, pages: list[dict]) -> None:
    rows = []
    for page in pages or []:
        if page.get("status") != "checked" or not page.get("result"):
            continue
        for check in page.get("result", {}).get("checks", []):
            if check.get("status") in {"warn", "fail"}:
                rows.append(
                    (
                        page.get("url", ""),
                        check.get("name", ""),
                        check.get("status", "").title(),
                        check.get("detail", ""),
                        check.get("recommendation", ""),
                    )
                )
    if rows:
        doc.add_heading("Warnings And Fix Items", level=2)
        _add_simple_table(doc, ("Page", "Check", "Status", "Detail", "Recommendation"), rows)


def _add_seo_checks(doc: Document, checks: list[dict]) -> None:
    rows = [
        (
            check.get("name", ""),
            check.get("status", "").title(),
            check.get("detail", ""),
            check.get("recommendation", ""),
        )
        for check in checks or []
    ]
    if rows:
        doc.add_heading("Checks", level=2)
        _add_simple_table(doc, ("Check", "Status", "Detail", "Recommendation"), rows)


def _add_seo_link_health(doc: Document, links: dict) -> None:
    links = links or {}
    issue_groups = (
        ("404 Links", links.get("not_found_links", []) or []),
        ("Unreachable Links", links.get("unreachable_links", []) or []),
        ("Other Broken HTTP Links", links.get("broken_links", []) or []),
        ("Redirecting Links", links.get("redirect_links", []) or []),
    )
    added_heading = False
    for heading, items in issue_groups:
        if not items:
            continue
        if not added_heading:
            doc.add_heading("Detailed Link Health", level=2)
            added_heading = True
        doc.add_heading(heading, level=3)
        rows = []
        for item in items:
            rows.append(
                (
                    item.get("source_page", ""),
                    item.get("url", ""),
                    item.get("status_code") or "n/a",
                    item.get("final_url", ""),
                    item.get("text", "") or "No anchor text",
                    item.get("type", ""),
                )
            )
        _add_simple_table(doc, ("Found On Page", "Link URL", "HTTP", "Final URL", "Anchor", "Type"), rows)
    if not added_heading:
        doc.add_heading("Detailed Link Health", level=2)
        doc.add_paragraph("No broken, unreachable, 404, or redirecting links were found in the sampled links.")


def _add_simple_table(doc: Document, headers: tuple[str, ...], rows: list[tuple] | list[list] | list[dict]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, label in enumerate(headers):
        table.rows[0].cells[index].text = str(label)
    for row_item in rows:
        cells = table.add_row().cells
        if isinstance(row_item, dict):
            values = [row_item.get(header, "") for header in headers]
        else:
            values = list(row_item)
        for index, value in enumerate(values[: len(headers)]):
            cells[index].text = str(value or "")


def _planner_items(plan: dict, key: str) -> list[dict]:
    items = plan.get(key, []) if isinstance(plan, dict) else []
    return [item for item in items if isinstance(item, dict)]


def _planner_keyword(name: str) -> str:
    cleaned = re.sub(r"\s+", " ", (name or "").replace("/", " ")).strip()
    return cleaned or "core page"


def _planner_page_keyword(page: dict, fallback: str = "core page") -> str:
    if not isinstance(page, dict):
        return _planner_keyword(fallback)
    return _planner_keyword(str(page.get("keyword") or page.get("name") or page.get("h1") or fallback))


def _planner_page_h1(page: dict, fallback: str = "core page") -> str:
    if not isinstance(page, dict):
        return _planner_keyword(fallback)
    return _planner_keyword(str(page.get("h1") or page.get("name") or page.get("keyword") or fallback))


def _planner_page_name(page: dict, fallback: str = "core page") -> str:
    if not isinstance(page, dict):
        return _planner_keyword(fallback)
    return _planner_keyword(str(page.get("name") or page.get("h1") or page.get("keyword") or fallback))


def _planner_page_slug(page: dict, fallback: str = "core page") -> str:
    if isinstance(page, dict):
        slug = str(page.get("slug", "") or "").strip()
        if slug:
            return slug
    return _planner_slug(fallback)


def _planner_heading_prompt(keyword: str, headings: list[str], target_market: str) -> str:
    if headings:
        return f"Use these extracted headings as the H2/H3 base for {keyword}: {', '.join(headings[:8])}."
    return f"Create H2s that explain how {keyword} works for {target_market or 'the target audience'}."


def _slug_text(value: str) -> str:
    cleaned = re.sub(r"[^a-z0-9]+", "-", (value or "").casefold()).strip("-")
    return cleaned or "page"


def _planner_slug(name: str) -> str:
    keyword = _planner_keyword(name)
    if keyword.casefold() in {"home", "homepage", "home page"}:
        return "/"
    return f"/{_slug_text(keyword)}/"


def _planner_secondary_keywords(keyword: str, client: str = "", target_market: str = "") -> list[str]:
    values = [f"{keyword} guide", f"best {keyword}", f"{keyword} tips"]
    if client:
        values.insert(0, f"{client} {keyword}")
    if target_market:
        values.append(f"{keyword} {target_market}")
    return _unique_text(values)[:5]


def _planner_keyword_rows(keyword: str, client: str, target_market: str) -> list[tuple[str, str, str, str, str]]:
    keywords = [keyword, *_planner_secondary_keywords(keyword, client, target_market)]
    return [(item, "Research", "TBD", "TBD", _planner_intent(item)) for item in _unique_text(keywords)[:6]]


def _planner_intent(keyword: str) -> str:
    value = keyword.casefold()
    if any(term in value for term in ("buy", "price", "bonus", "promo", "download", "login", "register")):
        return "Commercial"
    if any(term in value for term in ("how", "guide", "what", "tips")):
        return "Informational"
    return "Mixed"


def _planner_internal_link_rows(main_pages: list[dict], blogs: list[dict]) -> list[tuple[str, str, str, str]]:
    rows = []
    for page in main_pages:
        name = _planner_page_h1(page, _planner_page_keyword(page))
        keyword = _planner_page_keyword(page)
        rows.append(("Homepage", name, keyword, "Pass authority from the main hub to a core page"))
    for index, page in enumerate(main_pages):
        next_page = main_pages[(index + 1) % len(main_pages)] if main_pages else {}
        if next_page and next_page != page:
            rows.append((_planner_page_h1(page), _planner_page_h1(next_page), f"Learn more about {_planner_page_keyword(next_page)}", "Connect related core pages"))
    for index, blog in enumerate(blogs[:12]):
        target = _planner_target_page(main_pages, blog.get("name", ""), index)
        rows.append((blog.get("name", ""), _planner_page_h1(target, "Homepage"), _planner_page_keyword(target, "Homepage"), "Support the matching core page"))
    return rows or [("Homepage", "Core Pages", "main services", "Build the first internal links")]


def _planner_blog_rows(main_pages: list[dict], blogs: list[dict]) -> list[tuple[str, str, str, str, str]]:
    rows = []
    for index, blog in enumerate(blogs):
        topic = blog.get("name", f"Blog Topic {index + 1}")
        target = _planner_target_page(main_pages, topic, index)
        rows.append(
            (
                topic,
                blog.get("source", "Planner"),
                _planner_page_h1(target, "Homepage"),
                _planner_intent(topic),
                f"Write a helpful blog about {topic} and link naturally to {_planner_page_h1(target, 'the target page')}.",
            )
        )
    return rows or [("Supporting blog topic", "Planner", "Homepage", "Informational", "Create a supporting topic after core page keywords are finalized.")]


def _planner_core_mapping_rows(main_pages: list[dict], blogs: list[dict]) -> list[tuple[str, str, str, str, str]]:
    rows = []
    for index, page in enumerate(main_pages):
        name = _planner_page_h1(page, f"Core Page {index + 1}")
        keyword = _planner_page_keyword(page, name)
        linked_topics = [
            blog.get("name", "")
            for blog_index, blog in enumerate(blogs)
            if _planner_page_h1(_planner_target_page(main_pages, blog.get("name", ""), blog_index)) == name
        ][:3]
        rows.append((name, _planner_page_slug(page, keyword), keyword, "Core landing page", ", ".join(linked_topics) or "Add supporting blog topics"))
    return rows or [("Homepage", "/", "homepage", "Primary hub", "Add supporting blog topics")]


def _planner_launch_rows(main_pages: list[dict], trust_pages: list[dict], blogs: list[dict]) -> list[tuple[str, str, str, str]]:
    page_names = [_planner_page_h1(item, item.get("name", "")) for item in main_pages]
    trust_names = [item.get("name", "") for item in trust_pages]
    blog_names = [item.get("name", "") for item in blogs]
    return [
        ("Day 1", "Finalize homepage keyword, metadata, and hero content", page_names[0] if page_names else "Homepage", "Confirm main conversion path."),
        ("Day 2", "Draft top-priority core pages", ", ".join(page_names[1:3]) or "Core page drafts", "Use keyword clusters and content layout prompts."),
        ("Day 3", "Draft remaining core pages", ", ".join(page_names[3:6]) or "Remaining core pages", "Add internal links to homepage."),
        ("Day 4", "Publish trust and support pages", ", ".join(trust_names[:4]) or "Trust pages", "Trust pages are not keyword clusters, but support credibility."),
        ("Day 5", "Create first supporting blog topics", ", ".join(blog_names[:3]) or "First blog topics", "Link each blog to a mapped core page."),
        ("Day 6", "Add internal links and anchor text variations", "Internal link map", "Check every core page has inbound and outbound links."),
        ("Day 7", "QA metadata, indexability, tracking, and launch readiness", "Launch checklist", "Submit sitemap and monitor early performance."),
    ]


def _v2_detailed_keyword_rows(page: dict) -> list[tuple[str, str, str, str, str, str]]:
    rows = []
    primary = page.get("primary_keyword", "") or page.get("name", "")
    keyword_rows = page.get("keyword_rows", []) if isinstance(page.get("keyword_rows", []), list) else []
    seen = set()

    for item in keyword_rows:
        keyword = item.get("keyword", "")
        if not keyword:
            continue
        seen.add(keyword.casefold())
        rows.append(_v2_detailed_keyword_row(item))

    if primary and primary.casefold() not in seen:
        rows.insert(
            0,
            (
                primary,
                _metric_value(page.get("total_volume", "")),
                _metric_value(page.get("global_volume", "")),
                _metric_value(page.get("traffic_potential", "")),
                _metric_value(page.get("difficulty", "")),
                _clean_intents(page.get("intent", "")) or "Research",
            ),
        )

    for item in page.get("related_keywords", [])[:8]:
        if str(item).casefold() in seen:
            continue
        rows.append((item, "Research", "Research", "Research", "Research", _planner_intent(item)))

    return rows or [("Add keyword", "Research", "Research", "Research", "Research", "Research")]


def _v2_detailed_keyword_row(item: dict) -> tuple[str, str, str, str, str, str]:
    return (
        item.get("keyword", ""),
        _metric_value(item.get("volume", "")),
        _metric_value(item.get("global_volume", "")),
        _metric_value(item.get("traffic_potential", "")),
        _metric_value(item.get("difficulty", "")),
        _clean_intents(item.get("intents", "")) or _planner_intent(item.get("keyword", "")),
    )


def _metric_value(value) -> str:
    if value is None or value == "":
        return "Research"
    return str(value)


def _clean_intents(value: str) -> str:
    cleaned = str(value or "").replace("|", ",").replace(";", ",")
    parts = []
    seen = set()
    for part in re.split(r",|\n", cleaned):
        item = " ".join(part.strip().split())
        key = item.casefold()
        if item and key not in seen:
            seen.add(key)
            parts.append(item)
    return ", ".join(parts)


def _add_v2_detailed_page_layout(doc: Document, index: int, page: dict, target_market: str) -> None:
    page_name = page.get("name", f"Core Page {index}")
    keyword = page.get("primary_keyword", page_name)
    slug = page.get("slug", _planner_slug(keyword))
    h1 = page.get("h1", keyword)
    h2s = page.get("suggested_h2s", []) if isinstance(page.get("suggested_h2s", []), list) else []

    doc.add_heading(f"PAGE {index}: {page_name} {slug}", level=3)
    _add_label_value(doc, "Primary Keyword", keyword)
    if page.get("related_keywords"):
        _add_label_value(doc, "Secondary Keywords", ", ".join(page.get("related_keywords", [])[:8]))
    doc.add_paragraph("Content Sections")
    doc.add_paragraph(f"Breadcrumb: Home > {page_name}")

    doc.add_heading("Hero Section", level=4)
    _add_label_value(doc, "H1", h1)
    doc.add_paragraph(f"Hero introduction: introduce {page_name} for {target_market or 'the target audience'} using the main keyword '{keyword}' naturally.")
    doc.add_paragraph("CTA button: Use a short action phrase that matches the page intent.")

    doc.add_heading("Promotional Banner", level=4)
    doc.add_paragraph(f"Banner text: highlight the strongest benefit of {page_name}.")
    doc.add_paragraph("CTA: Use a clear next-step button.")

    doc.add_heading(f"{page_name} Introduction", level=4)
    doc.add_paragraph("Explain the page purpose, what users can do here, why it matters, and how it connects to the rest of the site.")

    doc.add_heading("Suggested H2 Sections", level=4)
    section_titles = h2s[:5] or [f"Why Choose {page_name}", "Key Features", "Mobile Experience", "Safety And Support"]
    for title in section_titles:
        doc.add_paragraph(f"H2: {title}")
        doc.add_paragraph(f"Section content: explain this angle with useful details, examples, proof points, and an internal link opportunity for {keyword}.")

    homepage_sections = page.get("homepage_sections", []) if isinstance(page.get("homepage_sections", []), list) else []
    if homepage_sections:
        doc.add_heading("Core Page Sections", level=4)
        for section in homepage_sections:
            doc.add_paragraph(f"H2: {section.get('h2', section.get('page_name', 'Core Page Section'))}")
            doc.add_paragraph(f"Target page: {section.get('target_slug', '')} [{section.get('page_name', '')}]")
            doc.add_paragraph(f"Primary keyword: {section.get('primary_keyword', '')}")
            doc.add_paragraph(section.get("summary", "Introduce the related core page and guide visitors to continue."))
            if section.get("cta"):
                doc.add_paragraph(f"CTA button: {section.get('cta')}")

    doc.add_heading("Frequently Asked Questions", level=4)
    doc.add_paragraph("H2: Frequently Asked Questions")
    for question in _v2_detailed_questions(page_name, keyword):
        doc.add_paragraph(question, style="List Bullet")

    doc.add_heading("Bottom CTA Banner", level=4)
    doc.add_paragraph(f"H2: Continue With {page_name}")
    doc.add_paragraph("Supporting paragraph: summarize the page value and guide readers toward the next action.")


def _v2_detailed_questions(page_name: str, keyword: str) -> list[str]:
    topic = page_name or keyword or "this page"
    return [
        f"What can users find on the {topic} page?",
        f"How does {keyword or topic} work?",
        f"Can users access {topic} on mobile?",
        f"What support options are available for {topic}?",
        "What should users check before getting started?",
    ]


def _v2_detailed_internal_link_rows(core_pages: list[dict], blogs: list[dict]) -> list[tuple[str, str, str, str]]:
    rows = []
    for page in core_pages:
        rows.append(("Homepage", page.get("name", ""), page.get("primary_keyword", page.get("name", "")), "Pass authority to a core category page"))
    for index, page in enumerate(core_pages):
        if len(core_pages) < 2:
            continue
        next_page = core_pages[(index + 1) % len(core_pages)]
        rows.append((page.get("name", ""), next_page.get("name", ""), f"Explore {next_page.get('primary_keyword', next_page.get('name', ''))}", "Connect related core page journeys"))
    for blog in blogs[:12]:
        rows.append((blog.get("name", ""), blog.get("target_page", "Homepage"), blog.get("primary_keyword", blog.get("name", "")), "Support the mapped core page"))
    return rows or [("Homepage", "Core Pages", "main pages", "Create the first internal links")]


def _add_v2_detailed_internal_link_map(doc: Document, core_pages: list[dict], support_pages: list[dict], blogs: list[dict]) -> None:
    target_pages = list(core_pages)
    if not target_pages:
        _add_simple_table(doc, ("Source Page", "Link To", "Suggested Anchor", "Purpose"), _v2_detailed_internal_link_rows(core_pages, blogs))
        return

    for index, page in enumerate(target_pages, start=1):
        page_name = page.get("name", f"Core Page {index}")
        keyword = page.get("primary_keyword", page_name)
        slug = page.get("slug", _planner_slug(keyword))

        doc.add_paragraph(f"Target page: {slug} [{page_name}]", style="Heading 4")
        doc.add_paragraph(f"Primary keyword: {keyword}")
        doc.add_paragraph(_v2_internal_link_description(page, keyword))
        doc.add_paragraph(f"CTA button: {_v2_internal_link_cta(page)}")

        if _v2_is_homepage(page):
            _add_simple_table(
                doc,
                ("Homepage Section", "Link To", "Suggested Anchor", "Purpose"),
                _v2_homepage_outbound_link_rows(core_pages, support_pages),
            )
            doc.add_paragraph("Homepage Blog Section", style="Heading 4")
            _add_simple_table(
                doc,
                ("Homepage Section", "Link To", "Suggested Anchor", "Purpose"),
                _v2_homepage_blog_section_rows(blogs),
            )
        else:
            _add_simple_table(
                doc,
                ("Linking Page", "Suggested Anchor", "Placement", "Purpose"),
                _v2_internal_link_source_rows(page, core_pages, blogs, minimum_links=3),
            )


def _v2_homepage_outbound_link_rows(core_pages: list[dict], support_pages: list[dict]) -> list[tuple[str, str, str, str]]:
    rows = []
    all_pages = [
        page
        for page in [*core_pages, *support_pages]
        if not _v2_is_homepage(page)
    ]
    for page in all_pages:
        page_name = page.get("name", "Page")
        keyword = page.get("primary_keyword", page_name)
        slug = page.get("slug", _planner_slug(keyword))
        rows.append(
            (
                f"{page_name} section",
                f"{slug} [{page_name}]",
                keyword,
                "Send homepage users to the matching core or support page.",
            )
        )
    return rows or [("Core pages section", "Add core page URLs", "main pages", "Link homepage to every non-blog page.")]


def _v2_homepage_blog_section_rows(blogs: list[dict]) -> list[tuple[str, str, str, str]]:
    category_names = _unique_text([blog.get("category", "") for blog in blogs if blog.get("category", "")])
    rows = [("Latest articles section", "/blog/ [Blog Hub]", "Read latest guides", "Create a homepage entry point to the blog hub.")]
    for category in category_names[:6]:
        rows.append(
            (
                f"{category} articles",
                f"/blog/category/{_slug_text(category)}/ [{category} Blog Category]",
                f"{category} guides",
                "Link to a blog category hub instead of individual blog posts.",
            )
        )
    return rows


def _v2_internal_link_source_rows(target_page: dict, core_pages: list[dict], blogs: list[dict], minimum_links: int = 3) -> list[tuple[str, str, str, str]]:
    target_name = target_page.get("name", "")
    target_keyword = target_page.get("primary_keyword", target_name)
    target_slug = target_page.get("slug", _planner_slug(target_keyword))
    source_rows = []
    seen_sources = set()

    for source in _v2_internal_link_sources(target_page, core_pages, blogs):
        source_name = source.get("name", "Homepage")
        source_slug = source.get("slug", _planner_slug(source_name))
        source_key = (source_slug or source_name).casefold()
        if not source_key or source_key in seen_sources:
            continue
        seen_sources.add(source_key)
        source_rows.append(
            (
                f"{source_slug} [{source_name}]",
                _v2_internal_anchor(target_keyword, source_name),
                _v2_internal_placement(source_name, target_name),
                f"Guide users to {target_slug} after they learn about {target_keyword}.",
            )
        )
        if len(source_rows) >= minimum_links:
            break

    fallback_sources = (
        ("Blog Hub", "/blog/"),
        ("Category Hub", f"/blog/category/{_slug_text(target_keyword)}/"),
        ("Related Guides Hub", f"/blog/tag/{_slug_text(target_keyword)}/"),
    )
    fallback_index = 0
    while len(source_rows) < minimum_links:
        fallback_name, fallback_slug = fallback_sources[min(fallback_index, len(fallback_sources) - 1)]
        fallback_index += 1
        source_rows.append(
            (
                f"{fallback_slug} [{fallback_name}]",
                _v2_internal_anchor(target_keyword, fallback_name),
                "Intro or related resources block",
                f"Add another crawlable path to {target_slug}.",
            )
        )

    return source_rows


def _v2_internal_link_sources(target_page: dict, core_pages: list[dict], blogs: list[dict]) -> list[dict]:
    target_name = (target_page.get("name", "") or "").casefold()
    target_slug = (target_page.get("slug", "") or "").casefold()
    sources = []

    homepage = next((page for page in core_pages if _v2_is_homepage(page)), None)
    if homepage:
        sources.append(homepage)

    for page in core_pages:
        if page is target_page:
            continue
        name = (page.get("name", "") or "").casefold()
        slug = (page.get("slug", "") or "").casefold()
        if name == target_name or slug == target_slug:
            continue
        sources.append(page)

    for blog in blogs:
        if (blog.get("target_page", "") or "").casefold() == target_name:
            sources.append(
                {
                    "name": blog.get("name", "Supporting Blog"),
                    "slug": f"/blog/{_slug_text(blog.get('name', 'supporting-blog'))}/",
                }
            )

    return sources


def _v2_internal_link_description(page: dict, keyword: str) -> str:
    page_name = page.get("name", keyword)
    intent = page.get("intent", "Mixed")
    minimum_label = "one related page" if _v2_is_homepage(page) else "at least three related pages"
    return (
        f"Use this page as the main destination for {keyword}. Add contextual links from {minimum_label} "
        f"so users and crawlers can reach the {page_name} cluster from multiple paths. Search intent: {intent}."
    )


def _v2_internal_link_cta(page: dict) -> str:
    page_name = page.get("name", "This Page")
    keyword = page.get("primary_keyword", page_name)
    if any(term in keyword.casefold() for term in ("login", "register", "signup", "sign up", "download")):
        return f"Go To {page_name}"
    if any(term in keyword.casefold() for term in ("bonus", "promo", "promotion")):
        return "View Current Promos"
    return f"Explore {page_name}"


def _v2_internal_anchor(keyword: str, source_name: str) -> str:
    if source_name.casefold() in {"homepage", "home", "home page"}:
        return keyword
    return f"Explore {keyword}"


def _v2_internal_placement(source_name: str, target_name: str) -> str:
    if source_name.casefold() in {"homepage", "home", "home page"}:
        return f"Homepage section for {target_name}"
    return "Relevant H2 section or bottom CTA block"


def _v2_is_homepage(page: dict) -> bool:
    name = (page.get("name", "") or "").strip().casefold()
    slug = (page.get("slug", "") or "").strip()
    keyword = (page.get("primary_keyword", "") or "").strip().casefold()
    return slug == "/" or name in {"home", "homepage", "home page"} or keyword in {"home", "homepage", "home page"}


def _v2_detailed_blog_rows(blogs: list[dict]) -> list[tuple]:
    rows = []
    for index, blog in enumerate(blogs, start=1):
        rows.append(
            (
                index,
                blog.get("name", f"Blog Topic {index}"),
                blog.get("primary_keyword", ""),
                blog.get("intent", _planner_intent(blog.get("name", ""))),
                blog.get("category", ""),
                blog.get("target_page", ""),
                blog.get("source", "Planner"),
            )
        )
    return rows or [(1, "Supporting blog topic", "category guide", "Informational", "General", "Homepage", "Planner")]


def _v2_detailed_blog_category_rows(category_plans: list[dict]) -> list[tuple]:
    rows = []
    for category in category_plans:
        rows.append(
            (
                category.get("name", ""),
                category.get("title", ""),
                f"{category.get('title_characters', '')}/{category.get('title_min_characters', '')}-{category.get('title_max_characters', '')}",
                category.get("meta_description", ""),
                f"{category.get('meta_description_characters', '')}/{category.get('meta_description_min_characters', '')}-{category.get('meta_description_max_characters', '')}",
                ", ".join(category.get("topics", [])[:5]) if isinstance(category.get("topics", []), list) else "",
            )
        )
    return rows or [("Category", "Add category title", "50-60", "Add a 130-150 character meta description.", "130-150", "Add topics")]


def _v2_detailed_core_mapping_rows(core_pages: list[dict], blogs: list[dict]) -> list[tuple[str, str, str, str, str]]:
    rows = []
    for page in core_pages:
        page_name = page.get("name", "")
        linked_blogs = [blog.get("name", "") for blog in blogs if blog.get("target_page") == page_name][:4]
        rows.append(
            (
                page_name,
                page.get("slug", _planner_slug(page.get("primary_keyword", page_name))),
                page.get("primary_keyword", ""),
                page.get("h1", ""),
                ", ".join(linked_blogs) or "Add supporting blog topics",
            )
        )
    return rows or [("Homepage", "/", "homepage", "Homepage", "Add supporting blog topics")]


def _v2_detailed_launch_rows(core_pages: list[dict], support_pages: list[dict], blogs: list[dict]) -> list[tuple[str, str, str, str]]:
    core_names = [page.get("name", "") for page in core_pages]
    support_names = [page.get("name", "") for page in support_pages]
    blog_names = [blog.get("name", "") for blog in blogs]
    return [
        ("Day 1", "Finalize homepage and top navigation", core_names[0] if core_names else "Homepage", "Use the homepage as the main silo hub"),
        ("Day 2", "Build highest-volume core page", core_names[1] if len(core_names) > 1 else "Core page", "Use keyword cluster table and page layout"),
        ("Day 3", "Build next core page", core_names[2] if len(core_names) > 2 else "Core page", "Add contextual links from homepage"),
        ("Day 4", "Publish support pages", ", ".join(support_names[:3]) or "Support pages", "Add footer/header trust links"),
        ("Day 5", "Publish supporting blog topics", ", ".join(blog_names[:2]) or "Blog topics", "Link each blog to its target core page"),
        ("Day 6", "Complete remaining priority core pages", ", ".join(core_names[3:6]) or "Remaining core pages", "Check slug, H1, H2, and internal anchors"),
        ("Day 7", "QA, metadata, indexing, and analytics", "Launch checklist", "Review noindex, canonicals, schema, and Search Console submission"),
    ]


def _planner_target_page(main_pages: list[dict], topic: str, index: int) -> dict:
    if not main_pages:
        return {"name": "Homepage"}
    topic_words = set(re.findall(r"[a-z0-9]+", (topic or "").casefold()))
    for page in main_pages:
        page_words = set(re.findall(r"[a-z0-9]+", " ".join([_planner_page_keyword(page), _planner_page_h1(page)]).casefold()))
        if topic_words & page_words:
            return page
    return main_pages[index % len(main_pages)]


def _unique_text(values: list[str]) -> list[str]:
    seen = set()
    output = []
    for value in values:
        cleaned = re.sub(r"\s+", " ", str(value or "")).strip()
        key = cleaned.casefold()
        if cleaned and key not in seen:
            seen.add(key)
            output.append(cleaned)
    return output


def _truncate_text(value: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}..."


def _add_gsc_totals_table(doc: Document, daily_rows: list[dict]) -> None:
    clicks = sum(float(row.get("clicks", 0) or 0) for row in daily_rows)
    impressions = sum(float(row.get("impressions", 0) or 0) for row in daily_rows)
    ctr = clicks / impressions if impressions else 0
    avg_position = _weighted_average_position(daily_rows)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for index, label in enumerate(("Clicks", "Impressions", "CTR", "Avg Position")):
        table.rows[0].cells[index].text = label
    row = table.add_row().cells
    row[0].text = f"{clicks:,.0f}"
    row[1].text = f"{impressions:,.0f}"
    row[2].text = f"{ctr:.2%}"
    row[3].text = f"{avg_position:.1f}" if avg_position else "n/a"


def _add_query_bar_table(doc: Document, rows: list[dict]) -> None:
    max_impressions = max((float(row.get("impressions", 0) or 0) for row in rows), default=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ("Query", "Clicks", "Impressions", "CTR", "Position", "Impression Graph")
    for index, label in enumerate(headers):
        table.rows[0].cells[index].text = label
    for item in rows:
        cells = table.add_row().cells
        cells[0].text = str(item.get("query", "") or "n/a")
        cells[1].text = f"{float(item.get('clicks', 0) or 0):,.0f}"
        cells[2].text = f"{float(item.get('impressions', 0) or 0):,.0f}"
        cells[3].text = f"{float(item.get('ctr', 0) or 0):.2%}"
        cells[4].text = f"{float(item.get('position', 0) or 0):.1f}"
        cells[5].text = _ascii_bar(float(item.get("impressions", 0) or 0), max_impressions)


def _add_section_table(doc: Document, rows: list[dict], keys: tuple[str, ...]) -> None:
    table = doc.add_table(rows=1, cols=len(keys))
    table.style = "Table Grid"
    for index, key in enumerate(keys):
        table.rows[0].cells[index].text = key.replace("_", " ").title()
    for item in rows:
        cells = table.add_row().cells
        for index, key in enumerate(keys):
            cells[index].text = str(item.get(key, ""))


def _gsc_report_sections(report: dict) -> list[tuple[str, list[dict], tuple[str, ...]]]:
    return [
        ("GSC Diagnosis", report.get("gsc_diagnosis", []), ("finding", "evidence", "meaning")),
        ("Opportunities", report.get("opportunities", []), ("opportunity", "reason", "recommended_action")),
        ("Recommendations", report.get("recommendations", []), ("priority", "area", "recommendation", "impact", "effort")),
        ("Backlink Analysis", report.get("backlink_analysis", []), ("finding", "evidence", "seo_effect", "recommended_action")),
        ("Content Plan", report.get("content_plan", []), ("section_or_asset", "target_query", "notes")),
        ("Technical Checks", report.get("technical_checks", []), ("check", "why", "how")),
        ("Monitoring Plan", report.get("monitoring_plan", []), ("metric", "target", "timing")),
        ("Next Steps", report.get("next_steps", []), ("step", "priority", "effort")),
    ]


def _ascii_bar(value: float, max_value: float, width: int = 24) -> str:
    if max_value <= 0:
        return ""
    filled = max(1, int(round((value / max_value) * width))) if value > 0 else 0
    return "#" * filled


def _weighted_average_position(rows: list[dict]) -> float:
    impressions = sum(float(row.get("impressions", 0) or 0) for row in rows)
    if not impressions:
        return 0
    return sum(float(row.get("position", 0) or 0) * float(row.get("impressions", 0) or 0) for row in rows) / impressions


def _download_filename(value: str) -> str:
    cleaned = "".join(character if character.isalnum() or character in {"-", "_"} else "_" for character in value.strip())
    cleaned = "_".join(part for part in cleaned.split("_") if part)
    return cleaned[:70] or "gsc_summary_report"


def _gsc_trend_png(rows: list[dict]) -> BytesIO | None:
    data = [row for row in rows if row.get("date")]
    if not data:
        return None
    width = 920
    height = 320
    padding = (58, 24, 34, 46)
    pixels = [[(255, 255, 255) for _ in range(width)] for _ in range(height)]
    left, top, right, bottom = padding
    plot_width = width - left - right
    plot_height = height - top - bottom
    _draw_line(pixels, left, top, left, top + plot_height, (220, 211, 199))
    _draw_line(pixels, left, top + plot_height, left + plot_width, top + plot_height, (220, 211, 199))
    _draw_series(pixels, data, "impressions", left, top, plot_width, plot_height, (166, 134, 87))
    _draw_series(pixels, data, "clicks", left, top, plot_width, plot_height, (77, 124, 15))
    return _encode_png(pixels)


def _draw_series(pixels, rows: list[dict], key: str, left: int, top: int, plot_width: int, plot_height: int, color: tuple[int, int, int]) -> None:
    max_value = max((float(row.get(key, 0) or 0) for row in rows), default=1) or 1
    points = []
    for index, row in enumerate(rows):
        x = left + (plot_width / 2 if len(rows) == 1 else (index / (len(rows) - 1)) * plot_width)
        y = top + plot_height - (float(row.get(key, 0) or 0) / max_value) * plot_height
        points.append((int(round(x)), int(round(y))))
    for start, end in zip(points, points[1:]):
        _draw_line(pixels, start[0], start[1], end[0], end[1], color, thickness=4)
    for x, y in points:
        _draw_circle(pixels, x, y, 4, color)


def _draw_line(pixels, x0: int, y0: int, x1: int, y1: int, color: tuple[int, int, int], thickness: int = 1) -> None:
    dx = x1 - x0
    dy = y1 - y0
    steps = max(abs(dx), abs(dy), 1)
    for step in range(steps + 1):
        x = int(round(x0 + dx * step / steps))
        y = int(round(y0 + dy * step / steps))
        radius = max(0, thickness // 2)
        for yy in range(y - radius, y + radius + 1):
            for xx in range(x - radius, x + radius + 1):
                _set_pixel(pixels, xx, yy, color)


def _draw_circle(pixels, cx: int, cy: int, radius: int, color: tuple[int, int, int]) -> None:
    for y in range(cy - radius, cy + radius + 1):
        for x in range(cx - radius, cx + radius + 1):
            if math.hypot(x - cx, y - cy) <= radius:
                _set_pixel(pixels, x, y, color)


def _set_pixel(pixels, x: int, y: int, color: tuple[int, int, int]) -> None:
    if 0 <= y < len(pixels) and 0 <= x < len(pixels[y]):
        pixels[y][x] = color


def _encode_png(pixels) -> BytesIO:
    height = len(pixels)
    width = len(pixels[0]) if height else 0
    raw = b"".join(b"\x00" + b"".join(bytes(pixel) for pixel in row) for row in pixels)
    png = b"\x89PNG\r\n\x1a\n"
    png += _png_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
    png += _png_chunk(b"IDAT", zlib.compress(raw, 9))
    png += _png_chunk(b"IEND", b"")
    output = BytesIO(png)
    output.seek(0)
    return output


def _png_chunk(kind: bytes, data: bytes) -> bytes:
    return struct.pack(">I", len(data)) + kind + data + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)

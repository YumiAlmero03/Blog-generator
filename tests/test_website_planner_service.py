import io
import json
import random
import re
import zipfile

from docx import Document

from app import create_app
from app.controllers import tool_controller
from app.services import website_planner_service
from app.services.website_planner_service import build_keyword_website_plan_v2, build_website_plan, enrich_keyword_website_plan_v2_with_ai, fetch_google_trends_blog_keywords, parse_ahrefs_keyword_csv, parse_keyword_categories, parse_page_list


def test_parse_page_list_deduplicates_lines_and_commas():
    pages = parse_page_list("Home\nAbout Us, Contact Us\nhome\n")

    assert pages == ["Home", "About Us", "Contact Us"]


def test_parse_keyword_categories_accepts_lists_and_deduplicates():
    categories = parse_keyword_categories(["Slots", "Sports, Casino Games", "slots"])

    assert categories == ["Slots", "Sports", "Casino Games"]


def test_build_website_plan_uses_requested_counts_with_random_lists():
    plan = build_website_plan(
        "Home\nAbout Us\nServices",
        "Privacy Policy\nTerms",
        page_count=4,
        trust_page_count=3,
        blog_count=2,
        keyword_categories_text=["Slots"],
        use_google_trends=False,
        rng=random.Random(7),
    )

    assert plan["summary"] == {
        "main_pages": 4,
        "trust_pages": 3,
        "blogs": 2,
        "total": 9,
        "blog_keyword_source": "Category fallback",
    }
    assert len(plan["main_pages"]) == 4
    assert len(plan["trust_pages"]) == 3
    assert len(plan["blogs"]) == 2
    assert all("slots" in item["name"] for item in plan["blogs"])


def test_build_website_plan_filters_brand_names_from_blog_keywords():
    plan = build_website_plan(
        "Home",
        "Privacy Policy",
        page_count=1,
        trust_page_count=1,
        blog_count=4,
        keyword_categories_text=["BrandX", "Slots"],
        brand_names=["BrandX"],
        use_google_trends=False,
        rng=random.Random(3),
    )

    blog_keywords = [item["name"] for item in plan["blogs"]]
    assert blog_keywords
    assert all("brandx" not in keyword.casefold() for keyword in blog_keywords)


def test_parse_ahrefs_keyword_csv_uses_keyword_column_and_deduplicates():
    rows = parse_ahrefs_keyword_csv(
        "Keyword,Volume,Global Volume,Traffic Potential,Keyword Difficulty,Parent Keyword,Intents,Category\n"
        "gperya,12000,18000,22000,20,gperya,\"Navigational, Branded\",Brand\n"
        "gperya,9000,12000,14000,20,gperya,Navigational,Brand\n"
        "gperya app,5000,7000,8500,12,gperya app,Transactional,App\n"
    )

    assert [row["keyword"] for row in rows] == ["gperya", "gperya app"]
    assert rows[0]["volume"] == 12000
    assert rows[0]["global_volume"] == 18000
    assert rows[0]["traffic_potential"] == 22000
    assert rows[0]["difficulty"] == "20"
    assert rows[0]["intents"] == "Navigational, Branded"
    assert rows[1]["category"] == "App"


def test_parse_ahrefs_keyword_csv_accepts_utf16_tab_export():
    csv_text = (
        '"#"\t"Keyword"\t"Country"\t"Difficulty"\t"Volume"\t"Parent Keyword"\t"Intents"\t"Category"\n'
        '"1"\t"free online games"\t"ph"\t"93"\t"5600"\t"crazy games"\t"Informational"\t"Browser games"\n'
        '"2"\t"play free online games"\t"ph"\t"91"\t"3800"\t"crazy games"\t"Informational"\t"Casual games"\n'
    )

    rows = parse_ahrefs_keyword_csv(csv_text.encode("utf-16"))

    assert [row["keyword"] for row in rows] == ["free online games", "play free online games"]
    assert rows[0]["volume"] == 5600
    assert rows[0]["difficulty"] == "93"
    assert rows[0]["parent_keyword"] == "crazy games"
    assert rows[1]["category"] == "Casual games"


def test_build_keyword_website_plan_v2_meets_minimum_plan_counts():
    rows = parse_ahrefs_keyword_csv(
        "Keyword,Volume,Difficulty,Parent Keyword,Intents,Category\n"
        "gperya,12000,20,gperya,Navigational,Brand\n"
        "gperya login,7000,10,gperya login,Navigational,Login\n"
        "gperya app,5000,12,gperya app,Transactional,App\n"
        "gperya casino,3000,18,gperya casino,Transactional,Casino\n"
        "gperya sports betting,2100,16,gperya sports betting,Transactional,Sports\n"
        "gperya slots,1900,14,gperya slots,Transactional,Slots\n"
        "gperya bonus,1500,9,gperya bonus,Commercial,Promotions\n"
        "is gperya legit,700,7,gperya review,Informational,Trust\n"
        "how to withdraw gperya,600,8,gperya withdrawal,Informational,Payments\n"
        "gperya customer service,400,5,gperya support,Support,Support\n"
    )

    plan = build_keyword_website_plan_v2(rows, source_filename="ahrefs.csv")

    assert plan["source_filename"] == "ahrefs.csv"
    assert plan["summary"]["core_page_count"] >= 9
    assert plan["summary"]["support_page_count"] >= 4
    assert plan["summary"]["blog_count"] >= 6
    assert plan["summary"]["category_count"] >= 6
    assert any(page["primary_keyword"] == "gperya app" for page in plan["core_pages"])
    assert plan["core_pages"][0]["name"] == "Homepage"
    assert plan["core_pages"][0]["primary_keyword"] == "gperya"
    assert plan["core_pages"][0]["homepage_sections"]
    assert any(section["page_name"] == "Sports Betting Page" for section in plan["core_pages"][0]["homepage_sections"])
    assert any(page["name"] == "Sports Betting Page" and page["primary_keyword"] == "gperya sports betting" for page in plan["core_pages"])
    assert any(blog["category"] for blog in plan["blogs"])
    assert plan["wrap_up"]


def test_build_keyword_website_plan_v2_keeps_homepage_as_page_one_with_high_volume_category():
    rows = parse_ahrefs_keyword_csv(
        "Keyword,Volume,Difficulty,Parent Keyword,Intents,Category\n"
        "sports betting gperya,50000,20,sports betting,Transactional,Sports\n"
        "gperya,12000,20,gperya,Navigational,General Keywords\n"
    )

    plan = build_keyword_website_plan_v2(rows, blog_count=50, category_suggestions="Esports\nBasketball")

    assert plan["core_pages"][0]["name"] == "Homepage"
    assert plan["core_pages"][0]["index"] == 1
    assert plan["core_pages"][0]["slug"] == "/"
    assert any(category["name"] == "Esports" for category in plan["categories"])
    assert not any(page["name"] == "Esports Page" for page in plan["core_pages"])
    assert any(blog["category"] == "Esports" for blog in plan["blogs"])
    assert any(blog["primary_keyword"] == "best esports" for blog in plan["blogs"])
    esports_plan = next(category for category in plan["blog_category_plans"] if category["name"] == "Esports")
    assert 50 <= esports_plan["title_characters"] <= 60
    assert 130 <= esports_plan["meta_description_characters"] <= 150
    assert esports_plan["blog_count"] >= 1


def test_build_keyword_website_plan_v2_uses_brand_topic_fallback_when_cluster_has_no_keyword():
    rows = parse_ahrefs_keyword_csv(
        "Keyword,Volume,Difficulty,Parent Keyword,Intents,Category\n"
        "gperya com,12000,20,gperya,Navigational,Brand\n"
        "gperya login,7000,10,gperya login,Navigational,Login\n"
        "gperya app,5000,12,gperya app,Transactional,App\n"
    )

    plan = build_keyword_website_plan_v2(rows, source_filename="ahrefs.csv")

    sports_page = next(page for page in plan["core_pages"] if page["name"] == "Sports Betting Hub")
    assert sports_page["primary_keyword"] == "sports betting"
    assert sports_page["slug"] == "/sports-betting/"
    assert plan["core_pages"][0]["slug"] == "/"


def test_build_keyword_website_plan_v2_can_use_google_trends_for_blogs(monkeypatch):
    monkeypatch.setattr(
        website_planner_service,
        "fetch_google_trends_blog_keywords",
        lambda categories, brand_names, geo="PH": [
            {"name": "slots bonus guide", "source": "Google Trends", "trend_url": "https://trends.google.com/example"}
        ],
    )
    rows = parse_ahrefs_keyword_csv(
        "Keyword,Volume,Difficulty,Parent Keyword,Intents,Category\n"
        "gperya,12000,20,gperya,Navigational,General Keywords\n"
        "gperya slots,1900,14,gperya slots,Transactional,Slots\n"
    )

    plan = build_keyword_website_plan_v2(rows, blog_count=6, use_google_trends=True, google_trends_geo="PH", brand_names=["gperya"])

    assert plan["summary"]["blog_keyword_source"] == "Google Trends + category fallback"
    assert plan["blogs"][0]["primary_keyword"] == "slots bonus guide"
    assert plan["blogs"][0]["source"] == "Google Trends"
    assert plan["blogs"][0]["trend_url"] == "https://trends.google.com/example"
    assert all(blog["source"] != "Ahrefs CSV" for blog in plan["blogs"])


def test_enrich_keyword_website_plan_v2_with_ai_updates_headings():
    class FakeProvider:
        def generate_json(self, prompt):
            assert "Create stronger H1" in prompt
            assert "blog category meta titles" in prompt
            return json.dumps(
                {
                    "pages": [
                        {
                            "index": 1,
                            "h1": "Play Gperya Casino Games Online",
                            "suggested_h2s": ["Game Library Highlights", "Mobile Play Options", "Bonuses For New Players", "Payment Methods"],
                        }
                    ],
                    "blog_categories": [
                        {
                            "index": 1,
                            "title": "Fresh Casino Strategy Stories For Filipino Players",
                            "meta_description": "Read practical casino blog topics with sharper tips, safer play advice, and timely ideas made for Filipino players comparing their options.",
                        }
                    ],
                }
            )

    plan = {
        "core_pages": [
            {
                "index": 1,
                "name": "Homepage",
                "primary_keyword": "gperya",
                "related_keywords": ["gperya login"],
                "intent": "Navigational",
                "h1": "Gperya",
                "suggested_h2s": ["Old H2"],
            }
        ],
        "blog_category_plans": [
            {
                "index": 1,
                "name": "Casino Games",
                "title": "Casino Games Blog Ideas And Updates useful reader",
                "title_characters": 49,
                "title_min_characters": 50,
                "title_max_characters": 60,
                "meta_description": "Explore Casino Games articles with practical tips, comparisons, updates, and beginner-friendly answers for readers planning their next step.",
                "meta_description_characters": 131,
                "meta_description_min_characters": 130,
                "meta_description_max_characters": 150,
                "blog_count": 2,
                "topics": ["Best Casino Games For Beginners"],
            }
        ],
    }

    enriched = enrich_keyword_website_plan_v2_with_ai(FakeProvider(), plan, {"client": "Gperya"})

    assert enriched["core_pages"][0]["h1"] == "Play Gperya Casino Games Online"
    assert enriched["core_pages"][0]["suggested_h2s"][0] == "Game Library Highlights"
    assert enriched["core_pages"][0]["heading_source"] == "AI"
    category_plan = enriched["blog_category_plans"][0]
    assert category_plan["title"].startswith("Fresh Casino Strategy Stories")
    assert 50 <= category_plan["title_characters"] <= 60
    assert category_plan["meta_description"].startswith("Read practical casino blog topics")
    assert 130 <= category_plan["meta_description_characters"] <= 150
    assert category_plan["seo_copy_source"] == "AI"


def test_fetch_google_trends_blog_keywords_filters_by_category_and_brand(monkeypatch):
    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self):
            return b"""
            <rss>
              <channel>
                <item><title>Slots Bonus Guide</title></item>
                <item><title>BrandX Casino Offer</title></item>
                <item><title>Weather Today</title></item>
              </channel>
            </rss>
            """

    monkeypatch.setattr(website_planner_service, "urlopen", lambda request, timeout: FakeResponse())

    keywords = fetch_google_trends_blog_keywords(["Slots"], ["BrandX"], geo="PH")

    assert [item["name"] for item in keywords] == ["slots bonus guide"]
    assert keywords[0]["source"] == "Google Trends"
    assert "trends.google.com" in keywords[0]["trend_url"]


def test_website_planner_page_renders_and_generates_plan(monkeypatch):
    monkeypatch.setattr(website_planner_service, "fetch_google_trends_blog_keywords", lambda *args, **kwargs: [])

    app = create_app()
    app.testing = True
    client = app.test_client()

    response = client.get("/website-planner")
    html = response.get_data(as_text=True)

    assert response.status_code == 200
    assert "Website Planner" in html
    assert "How many blogs" in html
    assert "How many pages" in html
    assert "How many trust pages" in html
    assert "Keyword Category" in html
    assert "Slots" in html

    response = client.post(
        "/website-planner",
        data={
            "page_count": "2",
            "trust_page_count": "1",
            "blog_count": "3",
            "keyword_categories": ["Slots"],
            "planner_reference_file": (_docx_upload(), "reference.docx"),
        },
    )
    html = response.get_data(as_text=True)

    assert response.status_code == 200
    assert "5 total item(s)" in html
    assert "VIP Slots Guide" in html
    assert "Headings: Weekly Promotions, Fast Deposits" in html
    assert "Extracted 1 page(s) from DOCX Heading 1/Page sections." in html
    assert "Using uploaded DOCX pages as the core page list." in html
    assert "slots" in html.lower()
    assert "Blog Keywords" in html
    assert "Google Trends" in html
    assert "Print Preview" in html
    assert "Report sections:" in html
    assert "Download report (.docx)" in html
    assert "Using reference: reference.docx" in html
    assert "VIP rewards and weekly slot promotions" in html


def test_website_planner_v2_page_renders_and_groups_uploaded_csv(monkeypatch):
    class FakeProvider:
        def generate_json(self, prompt):
            return json.dumps(
                {
                    "pages": [
                        {
                            "index": 1,
                            "h1": "Gperya Online Casino Philippines",
                            "suggested_h2s": ["Games And Betting Options", "Account Access", "App Download", "Promos And Payments"],
                        }
                    ]
                }
            )

    monkeypatch.setattr(tool_controller, "get_provider", lambda: FakeProvider())

    app = create_app()
    app.testing = True
    client = app.test_client()

    response = client.get("/website-planner-v2")
    html = response.get_data(as_text=True)

    assert response.status_code == 200
    assert "Website Planner V2" in html
    assert "Ahrefs CSV" in html
    assert "Blog Keyword Generator" in html
    assert "Category Suggestions" in html

    csv_text = (
        "Keyword,Volume,Difficulty,Parent Keyword,Intents,Category\n"
        "gperya,12000,20,gperya,Navigational,Brand\n"
        "gperya login,7000,10,gperya login,Navigational,Login\n"
        "gperya app,5000,12,gperya app,Transactional,App\n"
        "gperya casino,3000,18,gperya casino,Transactional,Casino\n"
        "gperya sports betting,2100,16,gperya sports betting,Transactional,Sports\n"
        "gperya slots,1900,14,gperya slots,Transactional,Slots\n"
        "gperya bonus,1500,9,gperya bonus,Commercial,Promotions\n"
        "is gperya legit,700,7,gperya review,Informational,Trust\n"
    )
    response = client.post(
        "/website-planner-v2",
        data={
            "planner_client": "Gperya",
            "blog_count": "8",
            "category_suggestions": "Esports\nBasketball",
            "ahrefs_csv_file": (io.BytesIO(csv_text.encode("utf-8")), "ahrefs.csv"),
        },
        content_type="multipart/form-data",
    )
    html = response.get_data(as_text=True)

    assert response.status_code == 200
    assert "8 keyword(s) grouped" in html
    assert "Core Page Clusters" in html
    assert "Keyword Categories" in html
    assert "Blog Plan" in html
    assert "8 Blogs" in html
    assert "Download report (.docx)" in html
    assert "Download detailed report (.docx)" in html
    assert "Esports" in html
    assert "Homepage Sections For Core Pages" in html
    assert "Blog Category Titles And Meta Descriptions" in html
    assert "Category Title" in html
    assert "Meta Description" in html
    assert "gperya app" in html
    assert "Grouped keywords into" in html


def test_website_planner_v2_download_report_contains_core_pages():
    app = create_app()
    app.testing = True
    client = app.test_client()
    plan = {
        "core_pages": [
            {
                "index": 1,
                "name": "Homepage",
                "primary_keyword": "gperya com",
                "slug": "/",
                "h1": "Gperya Online Casino Philippines",
                "suggested_h2s": ["Games And Betting Options", "Promos And Payments"],
            },
            {
                "index": 2,
                "name": "Sports Betting Hub",
                "primary_keyword": "sports betting",
                "slug": "/sports-betting/",
                "h1": "Sports Betting Guide",
                "suggested_h2s": ["Markets To Cover"],
            },
        ],
        "support_pages": [{"index": 1, "name": "About Us", "primary_keyword": "about us", "slug": "/about-us/", "h1": "About Us"}],
        "blogs": [{"index": 1, "name": "Slots Bonus Guide", "primary_keyword": "slots bonus guide", "source": "Google Trends"}],
        "summary": {"core_page_count": 2, "support_page_count": 1, "blog_count": 1, "blog_keyword_source": "Google Trends + category fallback"},
    }

    response = client.post(
        "/website-planner-v2/download-report",
        data={
            "planner_client": "Gperya",
            "planner_domain": "https://example.com",
            "planner_target_market": "Philippines",
            "planner_language": "English",
            "planner_site_type": "Informational / Affiliate Lead Gen",
            "source_filename": "ahrefs.csv",
            "plan_json": json.dumps(plan),
        },
    )

    assert response.status_code == 200
    assert response.headers["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    text = _docx_text(response.data)
    assert "Gperya Online Casino Philippines" in text
    assert "CLUSTER 1: Homepage" in text
    assert "Page 1: Homepage" in text
    assert "sports betting" in text
    assert "/sports-betting/" in text


def test_website_planner_v2_download_detailed_report_matches_reference_sections():
    app = create_app()
    app.testing = True
    client = app.test_client()
    plan = {
        "core_pages": [
            {
                "index": 1,
                "name": "Homepage",
                "primary_keyword": "gperya com",
                "slug": "/",
                "h1": "Gperya Online Casino Philippines",
                "suggested_h2s": ["Games And Betting Options", "Promos And Payments"],
                "related_keywords": ["gperya login", "gperya app"],
                "keyword_rows": [
                    {
                        "keyword": "gperya com",
                        "volume": 12000,
                        "global_volume": 18000,
                        "traffic_potential": 22000,
                        "difficulty": "20",
                        "intents": "Navigational, Branded",
                    },
                    {
                        "keyword": "gperya login",
                        "volume": 7000,
                        "global_volume": 9000,
                        "traffic_potential": 11000,
                        "difficulty": "10",
                        "intents": "Navigational",
                    },
                ],
                "homepage_sections": [
                    {
                        "page_name": "Sports Betting Hub",
                        "target_slug": "/sports-betting/",
                        "primary_keyword": "sports betting",
                        "h2": "Explore Sports Betting Hub",
                        "summary": "Introduce sports betting and link to the sports betting hub.",
                        "cta": "View Sports Betting Hub",
                    }
                ],
                "total_volume": 12000,
                "intent": "Navigational",
            },
            {
                "index": 2,
                "name": "Sports Betting Hub",
                "primary_keyword": "sports betting",
                "slug": "/sports-betting/",
                "h1": "Sports Betting Guide",
                "suggested_h2s": ["Markets To Cover"],
                "related_keywords": [],
                "keyword_rows": [
                    {
                        "keyword": "sports betting",
                        "volume": 2100,
                        "global_volume": 4200,
                        "traffic_potential": 6200,
                        "difficulty": "16",
                        "intents": "Commercial, Transactional",
                    }
                ],
                "total_volume": 0,
                "intent": "Transactional",
            },
        ],
        "support_pages": [{"index": 1, "name": "About Us", "primary_keyword": "about us", "slug": "/about-us/", "h1": "About Us"}],
        "blogs": [
            {
                "index": 1,
                "name": "Slots Bonus Guide",
                "primary_keyword": "slots bonus guide",
                "source": "Google Trends",
                "category": "Slots",
                "target_page": "Homepage",
                "intent": "Informational",
            }
        ],
        "categories": [{"index": 1, "name": "Slots", "keyword_count": 1, "total_volume": 1900}],
        "blog_category_plans": [
            {
                "index": 1,
                "name": "Slots",
                "title": "Slots Blog Ideas And Updates useful reader guide",
                "title_characters": 49,
                "title_min_characters": 50,
                "title_max_characters": 60,
                "meta_description": "Explore Slots articles with practical tips, comparisons, updates, and beginner-friendly answers for readers planning their next step.",
                "meta_description_characters": 126,
                "meta_description_min_characters": 130,
                "meta_description_max_characters": 150,
                "blog_count": 1,
                "topics": ["Slots Bonus Guide"],
            }
        ],
        "summary": {"core_page_count": 2, "support_page_count": 1, "blog_count": 1, "category_count": 1, "total_volume": 12000, "blog_keyword_source": "Google Trends + category fallback"},
    }

    response = client.post(
        "/website-planner-v2/download-detailed-report",
        data={
            "planner_client": "Gperya",
            "planner_domain": "https://example.com",
            "planner_target_market": "Philippines",
            "planner_language": "English",
            "planner_site_type": "Informational / Affiliate Lead Gen",
            "source_filename": "ahrefs.csv",
            "plan_json": json.dumps(plan),
        },
    )

    assert response.status_code == 200
    assert response.headers["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "Website_Planner_V2_Detailed_Report" in response.headers["Content-Disposition"]
    text = _docx_text(response.data)
    for expected in (
        "SEO Master Plan - Gperya",
        "PART 1 - KEYWORD PLAN",
        "CLUSTER 1: Homepage",
        "Monthly Volume",
        "Global Volume",
        "Traffic Potential",
        "Keyword Difficulty",
        "Intents",
        "18000",
        "22000",
        "Navigational, Branded",
        "Commercial, Transactional",
        "PART 2 - CONTENT LAYOUT PLAN",
        "Site Architecture",
        "PAGE 2: Sports Betting Hub /sports-betting/",
        "Hero Section",
        "Promotional Banner",
        "Suggested H2 Sections",
        "H2: Markets To Cover",
        "Core Page Sections",
        "Explore Sports Betting Hub",
        "Frequently Asked Questions",
        "PART 3 - INTERNAL LINKING PLAN",
        "Target page: / [Homepage]",
        "Primary keyword: gperya com",
        "Homepage Blog Section",
        "/about-us/ [About Us]",
        "/blog/category/slots/ [Slots Blog Category]",
        "Target page: /sports-betting/ [Sports Betting Hub]",
        "Primary keyword: sports betting",
        "CTA button: Explore Sports Betting Hub",
        "/ [Homepage]",
        "/blog/ [Blog Hub]",
        "/blog/category/sports-betting/ [Category Hub]",
        "Anchor Text Guidelines",
        "PART 4 - BLOG CONTENT CATEGORIES AND CURRENT PAGE-DERIVED TOPICS",
        "Blog Category Title And Meta Description Plan",
        "PART 5 - SUMMARY DASHBOARD",
        "Current Core Page Mapping",
        "PART 6 - PRIORITY LAUNCH ORDER",
    ):
        assert expected in text


def test_website_planner_download_report_contains_required_sections():
    app = create_app()
    app.testing = True
    client = app.test_client()
    plan = {
        "main_pages": [
            {"index": 1, "name": "online slot games", "keyword": "online slot games", "h1": "Best Online Slot Games", "headings": ["Weekly Promotions", "Fast Deposits"]},
            {"index": 2, "name": "Home"},
        ],
        "trust_pages": [{"index": 1, "name": "Privacy Policy"}],
        "blogs": [{"index": 1, "name": "slot bonus guide", "source": "Planner"}],
        "summary": {"main_pages": 2, "trust_pages": 1, "blogs": 1, "total": 4, "blog_keyword_source": "Category fallback"},
    }

    response = client.post(
        "/website-planner/download-report",
        data={
            "planner_client": "PhMacao Casino",
            "planner_domain": "https://example.com",
            "planner_target_market": "Philippines",
            "planner_language": "English",
            "planner_site_type": "Informational / Affiliate Lead Gen",
            "planner_reference_filename": "reference.docx",
            "planner_reference_content": "Existing website content mentions VIP rewards, weekly slot promotions, and fast deposits.",
            "plan_json": json.dumps(plan),
        },
    )

    assert response.status_code == 200
    assert response.headers["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "Website_Planner_Report" in response.headers["Content-Disposition"]
    text = _docx_text(response.data)
    for expected in (
        "KEYWORD PLAN",
        "About site",
        "Keyword CLUSTER",
        "CONTENT LAYOUT PLAN",
        "INTERNAL LINKING PLAN",
        "Internal Link Map",
        "Anchor Text Guidelines",
        "BLOG CONTENT CATEGORIES AND CURRENT PAGE-DERIVED TOPICS",
        "SUMMARY DASHBOARD",
        "Current Core Page Mapping",
        "PRIORITY LAUNCH ORDER (First 7 Days)",
        "Step-by-step prompt",
        "Website Content Reference",
        "VIP rewards",
        "Best Online Slot Games",
        "Weekly Promotions",
    ):
        assert expected in text
    assert "Privacy PolicyPrimary Target Page" not in text


def test_settings_page_includes_website_planner_lists():
    app = create_app()
    app.testing = True
    response = app.test_client().get("/settings")

    html = response.get_data(as_text=True)
    assert response.status_code == 200
    assert "Website Planner Pages" in html
    assert "Pages Main" in html
    assert "Trust Page" in html


def _docx_text(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    return re.sub(r"<[^>]+>", "", xml)


def _docx_upload() -> io.BytesIO:
    doc = Document()
    doc.add_paragraph("VIP Slots Guide", style="Heading 1")
    doc.add_paragraph("Keyword: vip slots")
    doc.add_paragraph("Weekly Promotions", style="Heading 2")
    doc.add_paragraph("VIP rewards and weekly slot promotions.")
    doc.add_paragraph("Fast Deposits", style="Heading 2")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

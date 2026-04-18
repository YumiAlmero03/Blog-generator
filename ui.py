from flask import Flask, render_template, request, make_response, send_from_directory, url_for
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from pathlib import Path
from uuid import uuid4
from werkzeug.utils import secure_filename

from config import PROVIDER, MODEL
from database import (
    check_keyword_usage,
    get_brand_context,
    get_brand_record,
    list_brand_names,
    list_brand_records,
    record_blog,
    record_page,
    upsert_brand,
)
from generators.title_generator import generate_titles
from generators.meta_description_generator import generate_meta_descriptions
from generators.content_generator import generate_content
from generators.page_generator import generate_page
from generators.simple_page_generator import generate_simple_page
from logger import logger
from providers.base import ProviderError

app = Flask(__name__)

UPLOAD_ROOT = Path(__file__).resolve().parent / "data" / "uploads"
BRAND_LOGO_DIR = UPLOAD_ROOT / "brand_logos"
IMAGE_TOOL_DIR = UPLOAD_ROOT / "image_tools"
ALLOWED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}

for directory in (BRAND_LOGO_DIR, IMAGE_TOOL_DIR):
    directory.mkdir(parents=True, exist_ok=True)


def generation_error_message(default_message: str, exc: Exception) -> str:
    if isinstance(exc, ProviderError):
        return str(exc)
    return default_message


def allowed_image_file(filename: str) -> bool:
    suffix = Path(filename or "").suffix.lower()
    return suffix in ALLOWED_IMAGE_EXTENSIONS


def save_uploaded_image(file_storage, destination_dir: Path, prefix: str) -> str:
    filename = secure_filename(file_storage.filename or "")
    if not filename or not allowed_image_file(filename):
        raise ValueError("Please upload a PNG, JPG, JPEG, or WEBP image.")

    suffix = Path(filename).suffix.lower() or ".png"
    saved_name = f"{prefix}_{uuid4().hex}{suffix}"
    output_path = destination_dir / saved_name
    file_storage.save(output_path)
    return saved_name


def image_url(relative_path: str) -> str:
    cleaned = (relative_path or "").strip().replace("\\", "/")
    if not cleaned:
        return ""
    return url_for("uploaded_file", filename=cleaned)


def parse_ratio_dimensions(snap_ratio: str, fallback_width: int, fallback_height: int) -> tuple[int, int]:
    ratio_map = {
        "1:1": (1, 1),
        "4:5": (4, 5),
        "16:9": (16, 9),
        "9:16": (9, 16),
        "3:2": (3, 2),
        "original": (fallback_width, fallback_height),
    }
    ratio = ratio_map.get(snap_ratio)
    if not ratio or ratio[0] <= 0 or ratio[1] <= 0:
        return fallback_width, fallback_height
    return ratio


def calculate_output_dimensions(pixel_width: str, pixel_height: str, snap_ratio: str, fallback_width: int, fallback_height: int) -> tuple[int, int]:
    width_value = int(str(pixel_width or "").strip() or 0)
    height_value = int(str(pixel_height or "").strip() or 0)
    ratio_width, ratio_height = parse_ratio_dimensions(snap_ratio, fallback_width, fallback_height)

    if width_value <= 0 and height_value <= 0:
        raise ValueError("Enter a width or height for the final exported image.")
    if width_value > 0:
        return width_value, max(1, round(width_value * (ratio_height / ratio_width)))
    return max(1, round(height_value * (ratio_width / ratio_height))), height_value


def crop_image_to_box(image, crop_x: str, crop_y: str, crop_width: str, crop_height: str):
    crop_width_value = int(float(crop_width or 0))
    crop_height_value = int(float(crop_height or 0))
    if crop_width_value <= 0 or crop_height_value <= 0:
        raise ValueError("Set a crop area before processing the image.")
    if crop_width_value > image.width or crop_height_value > image.height:
        raise ValueError("The crop size is larger than the source image. Reduce the target crop size.")

    crop_x_value = int(float(crop_x or 0))
    crop_y_value = int(float(crop_y or 0))
    max_x = image.width - crop_width_value
    max_y = image.height - crop_height_value
    crop_x_value = min(max(crop_x_value, 0), max_x)
    crop_y_value = min(max(crop_y_value, 0), max_y)

    return image.crop(
        (
            crop_x_value,
            crop_y_value,
            crop_x_value + crop_width_value,
            crop_y_value + crop_height_value,
        )
    )


def snap_image_to_ratio(image, snap_ratio: str):
    if snap_ratio == "original":
        return image

    ratio_map = {
        "1:1": (1, 1),
        "4:5": (4, 5),
        "16:9": (16, 9),
        "9:16": (9, 16),
        "3:2": (3, 2),
    }
    target = ratio_map.get(snap_ratio)
    if not target:
        return image

    target_ratio = target[0] / target[1]
    width, height = image.size
    current_ratio = width / height

    if abs(current_ratio - target_ratio) < 0.0001:
        return image

    if current_ratio > target_ratio:
        new_width = round(height * target_ratio)
        left = max(0, (width - new_width) // 2)
        return image.crop((left, 0, left + new_width, height))

    new_height = round(width / target_ratio)
    top = max(0, (height - new_height) // 2)
    return image.crop((0, top, width, top + new_height))


def apply_logo_watermark(
    base_image,
    logo_image,
    position: str,
    opacity_percent: str,
    logo_scale_percent: str,
    watermark_x_percent: str = "",
    watermark_y_percent: str = "",
    watermark_rotation: str = "0",
):
    from PIL import Image

    base = base_image.convert("RGBA")
    logo = logo_image.convert("RGBA")

    opacity = max(0, min(100, int(opacity_percent or 45)))
    scale = max(1, int(logo_scale_percent or 20))

    target_logo_width = max(1, round(base.width * (scale / 100.0)))
    resize_ratio = target_logo_width / logo.width
    target_logo_height = max(1, round(logo.height * resize_ratio))
    logo = logo.resize((target_logo_width, target_logo_height), resample=Image.Resampling.LANCZOS)

    alpha = logo.getchannel("A")
    alpha = alpha.point(lambda value: round(value * (opacity / 100.0)))
    logo.putalpha(alpha)

    padding = max(16, round(min(base.width, base.height) * 0.03))

    rotation = float(str(watermark_rotation or "0").strip() or 0)
    if rotation:
        logo = logo.rotate(-rotation, expand=True, resample=Image.Resampling.BICUBIC)

    if str(watermark_x_percent).strip() and str(watermark_y_percent).strip():
        center_x = round(base.width * (float(watermark_x_percent) / 100.0))
        center_y = round(base.height * (float(watermark_y_percent) / 100.0))
        x = center_x - (logo.width // 2)
        y = center_y - (logo.height // 2)
    else:
        positions = {
            "top-left": (padding, padding),
            "top-right": (base.width - logo.width - padding, padding),
            "bottom-left": (padding, base.height - logo.height - padding),
            "center": ((base.width - logo.width) // 2, (base.height - logo.height) // 2),
            "bottom-right": (base.width - logo.width - padding, base.height - logo.height - padding),
        }
        x, y = positions.get(position, positions["bottom-right"])

    x = min(max(0, x), max(0, base.width - logo.width))
    y = min(max(0, y), max(0, base.height - logo.height))
    overlay = base.copy()
    overlay.alpha_composite(logo, (max(0, x), max(0, y)))
    return overlay


def html_to_docx_paragraph(doc, html_content):
    """Convert HTML content to docx paragraphs with formatting"""
    # Remove extra whitespace
    html_content = html_content.strip()
    if not html_content:
        return
    
    # Split by common HTML tags and process
    # Handle h2, h3, p, strong, br, a, ul, li tags
    lines = html_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Handle h2 headings
        if '<h2>' in line:
            text = re.sub(r'</?h2>', '', line)
            p = doc.add_paragraph(text.strip(), style='Heading 2')
            continue
        
        # Handle h3 headings
        if '<h3>' in line:
            text = re.sub(r'</?h3>', '', line)
            p = doc.add_paragraph(text.strip(), style='Heading 3')
            continue
        
        # Handle paragraphs and other content
        if html_content or line:
            # Parse inline formatting (bold, links, etc)
            p = doc.add_paragraph()
            parse_inline_html(p, line)

def parse_inline_html(paragraph, html_text):
    """Parse inline HTML and add formatted text to paragraph"""
    if not html_text:
        return
    
    # Replace <br> with newlines
    html_text = html_text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    
    # Split by tags to preserve formatting
    parts = re.split(r'(<a\s+href=["\']([^"\']+)["\']>|</a>|<strong>|</strong>|<b>|</b>)', html_text)
    
    current_bold = False
    current_url = None
    
    i = 0
    while i < len(parts):
        part = parts[i]
        
        if not part:
            i += 1
            continue
        
        # Check for link start
        if part.startswith('<a'):
            current_url = parts[i + 1] if i + 1 < len(parts) else None
            i += 2
            continue
        
        # Check for link end
        if part == '</a>':
            current_url = None
            i += 1
            continue
        
        # Check for bold tags
        if part in ['<strong>', '<b>']:
            current_bold = True
            i += 1
            continue
        
        if part in ['</strong>', '</b>']:
            current_bold = False
            i += 1
            continue
        
        # Regular text content
        if part and not part.startswith('<'):
            run = paragraph.add_run(part)
            if current_bold:
                run.bold = True
            if current_url:
                # Add hyperlink
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                r = run._element
                rPr = r.get_or_add_rPr()
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rPr.append(rStyle)
                # Create hyperlink element
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                r.addprevious(fldChar1)
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = f'HYPERLINK "{current_url}"'
                r.addprevious(instrText)
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                r.addnext(fldChar2)
        
        i += 1


def get_provider():
    if PROVIDER == "ollama":
        from providers.ollama_provider import OllamaProvider
        return OllamaProvider(MODEL)
    elif PROVIDER == "openai":
        from providers.openai_provider import OpenAIProvider
        return OpenAIProvider(MODEL)
    elif PROVIDER == "gemini":
        from providers.gemini_provider import GeminiProvider
        return GeminiProvider(MODEL)
    else:
        raise ValueError(f"Unsupported provider: {PROVIDER}")


@app.route("/", methods=["GET", "POST"])
def index():
    keyword = ""
    brand = ""
    supporting_keyword = ""
    tone = "natural"
    count = 10
    titles = []
    selected_title = ""
    meta_descriptions = []
    meta_description = ""
    content = ""
    error = None
    step = "title"  # "title", "meta_description", or "content"
    brand_names = list_brand_names()

    if request.method == "POST":
        action = request.form.get("action", "").strip()
        
        if action == "generate_titles":
            keyword = request.form.get("keyword", "").strip()
            brand = request.form.get("brand", "").strip()
            supporting_keyword = request.form.get("supporting_keyword", "").strip()
            tone = request.form.get("tone", "natural").strip() or "natural"
            count_raw = request.form.get("count", "10").strip()

            if not keyword:
                error = "Please enter one or more keywords."
            else:
                if brand:
                    upsert_brand(brand)
                try:
                    count = int(count_raw)
                except ValueError:
                    count = 10

                try:
                    provider = get_provider()
                    brand_context = get_brand_context(brand)
                    titles = generate_titles(
                        provider,
                        keyword=keyword,
                        tone=tone,
                        count=count,
                        brand=brand,
                        brand_context=brand_context,
                    )
                    step = "title"
                except Exception as exc:
                    logger.exception("generate_titles action failed")
                    error = generation_error_message(
                        "An error occurred while generating titles. Check logs/app.log for details.",
                        exc,
                    )
        
        elif action == "generate_meta":
            selected_title = request.form.get("selected_title", "").strip()
            keyword = request.form.get("keyword", "").strip()
            brand = request.form.get("brand", "").strip()
            supporting_keyword = request.form.get("supporting_keyword", "").strip()
            titles_raw = request.form.get("titles_json", "").strip()
            
            if not selected_title:
                error = "Please select a title first."
            else:
                try:
                    import json
                    titles = json.loads(titles_raw) if titles_raw else []
                    provider = get_provider()
                    if brand:
                        upsert_brand(brand)
                    brand_context = get_brand_context(brand)
                    meta_descriptions = generate_meta_descriptions(
                        provider,
                        title=selected_title,
                        keyword=keyword,
                        count=3,
                        brand=brand,
                        brand_context=brand_context,
                    )
                    if meta_descriptions:
                        meta_description = meta_descriptions[0]["text"]
                    step = "meta_description"
                except Exception as exc:
                    logger.exception("generate_meta action failed")
                    error = generation_error_message(
                        "An error occurred while generating meta descriptions. Check logs/app.log for details.",
                        exc,
                    )
        
        elif action == "generate_content":
            selected_title = request.form.get("selected_title", "").strip()
            keyword = request.form.get("keyword", "").strip()
            brand = request.form.get("brand", "").strip()
            supporting_keyword = request.form.get("supporting_keyword", "").strip()
            tone = request.form.get("tone", "natural").strip() or "natural"
            titles_raw = request.form.get("titles_json", "").strip()
            meta_descriptions_raw = request.form.get("meta_descriptions_json", "").strip()
            
            # Extract links from form
            links = []
            link_texts = request.form.getlist("link_text[]")
            link_urls = request.form.getlist("link_url[]")
            for text, url in zip(link_texts, link_urls):
                text = text.strip()
                url = url.strip()
                if text and url:
                    links.append({"text": text, "url": url})
            
            if not selected_title:
                error = "Please select a title first."
            else:
                try:
                    import json
                    titles = json.loads(titles_raw) if titles_raw else []
                    meta_descriptions = json.loads(meta_descriptions_raw) if meta_descriptions_raw else []
                    if meta_descriptions and not meta_description:
                        meta_description = meta_descriptions[0].get("text", "")
                    provider = get_provider()
                    if brand:
                        upsert_brand(brand)
                    brand_context = get_brand_context(brand)
                    content = generate_content(
                        provider,
                        title=selected_title,
                        keyword=keyword,
                        supporting_keyword=supporting_keyword,
                        tone=tone,
                        links=links,
                        brand=brand,
                        brand_context=brand_context,
                    )
                    record_blog(
                        brand=brand,
                        title=selected_title,
                        keyword=keyword,
                        supporting_keyword=supporting_keyword,
                    )
                    step = "content"
                except Exception as exc:
                    logger.exception("generate_content action failed")
                    error = generation_error_message(
                        "An error occurred while generating article content. Check logs/app.log for details.",
                        exc,
                    )

    return render_template(
        "index.html",
        provider=PROVIDER,
        model=MODEL,
        brand_names=brand_names,
        keyword=keyword,
        brand=brand,
        supporting_keyword=supporting_keyword,
        tone=tone,
        count=count,
        titles=titles,
        selected_title=selected_title,
        meta_descriptions=meta_descriptions,
        meta_description=meta_description,
        content=content,
        error=error,
        step=step,
    )


@app.route("/page-generator", methods=["GET", "POST"])
def page_generator():
    keyword = ""
    brand = ""
    supporting_keywords = ""
    page_type = ""
    expectations = ""
    page_title = ""
    meta_description = ""
    page_content = ""
    image_count = 0
    error = None
    brand_names = list_brand_names()

    if request.method == "POST":
        keyword = request.form.get("keyword", "").strip()
        brand = request.form.get("brand", "").strip()
        supporting_keywords = request.form.get("supporting_keywords", "").strip()
        page_type = request.form.get("page_type", "").strip()
        expectations = request.form.get("expectations", "").strip()

        if not keyword:
            error = "Please enter a keyword."
        else:
            try:
                provider = get_provider()
                if brand:
                    upsert_brand(brand)
                brand_context = get_brand_context(brand)
                result = generate_page(
                    provider,
                    keyword=keyword,
                    brand=brand,
                    supporting_keywords=supporting_keywords,
                    page_type=page_type,
                    expectations=expectations,
                    brand_context=brand_context,
                )
                page_title = result.get("title", "")
                meta_description = result.get("meta_description", "")
                page_content = result.get("content", "")
                image_count = result.get("image_count", 0)
                record_page(
                    brand=brand,
                    keyword=keyword,
                    page_title=page_title,
                    page_type=page_type,
                    supporting_keywords=supporting_keywords,
                    expectations=expectations,
                )
            except Exception as exc:
                logger.exception("page_generator action failed")
                error = generation_error_message(
                    "An error occurred while generating the page. Check logs/app.log for details.",
                    exc,
                )

    return render_template(
        "page_generator.html",
        provider=PROVIDER,
        model=MODEL,
        brand_names=brand_names,
        keyword=keyword,
        brand=brand,
        supporting_keywords=supporting_keywords,
        page_type=page_type,
        expectations=expectations,
        page_title=page_title,
        meta_description=meta_description,
        page_content=page_content,
        image_count=image_count,
        error=error,
    )


@app.route("/simple-page-generator", methods=["GET", "POST"])
def simple_page_generator():
    brand = ""
    page_title = ""
    page_type = ""
    expectations = ""
    generated_title = ""
    generated_content = ""
    error = None
    brand_names = list_brand_names()

    if request.method == "POST":
        brand = request.form.get("brand", "").strip()
        page_title = request.form.get("page_title", "").strip()
        page_type = request.form.get("page_type", "").strip()
        expectations = request.form.get("expectations", "").strip()

        if not page_title:
            error = "Please enter the page title or page name."
        else:
            try:
                provider = get_provider()
                if brand:
                    upsert_brand(brand)
                brand_context = get_brand_context(brand)
                result = generate_simple_page(
                    provider,
                    page_title=page_title,
                    page_type=page_type,
                    brand=brand,
                    expectations=expectations,
                    brand_context=brand_context,
                )
                generated_title = result.get("title", "")
                generated_content = result.get("content", "")
                record_page(
                    brand=brand,
                    keyword=page_title,
                    page_title=generated_title or page_title,
                    page_type=page_type or "simple page",
                    supporting_keywords="",
                    expectations=expectations,
                )
            except Exception as exc:
                logger.exception("simple_page_generator action failed")
                error = generation_error_message(
                    "An error occurred while generating the simple page. Check logs/app.log for details.",
                    exc,
                )

    return render_template(
        "simple_page_generator.html",
        provider=PROVIDER,
        model=MODEL,
        brand_names=brand_names,
        brand=brand,
        page_title=page_title,
        page_type=page_type,
        expectations=expectations,
        generated_title=generated_title,
        generated_content=generated_content,
        error=error,
    )


@app.route("/uploads/<path:filename>")
def uploaded_file(filename):
    return send_from_directory(UPLOAD_ROOT, filename)


@app.route("/text-tools")
def text_tools():
    return render_template(
        "text_tools.html",
        provider=PROVIDER,
        model=MODEL,
    )


@app.route("/image-tools", methods=["GET", "POST"])
def image_tools():
    brand = ""
    brand_logo_url = ""
    source_image_url = ""
    source_image_name = ""
    result_image_url = ""
    result_download_name = ""
    error = None
    success = None
    pixel_width = "800"
    pixel_height = "450"
    snap_ratio = "16:9"
    watermark_position = "bottom-right"
    watermark_opacity = "100"
    logo_scale = "20"
    output_filename = "watermarked-image"
    output_format = "webp"
    crop_x = "0"
    crop_y = "0"
    crop_width = ""
    crop_height = ""
    crop_scale = "70"
    watermark_x = "85"
    watermark_y = "85"
    watermark_rotation = "0"
    use_watermark = True
    brand_names = list_brand_names()

    if request.method == "POST":
        brand = request.form.get("brand", "").strip()
        pixel_width = request.form.get("pixel_width", "").strip()
        pixel_height = request.form.get("pixel_height", "").strip()
        snap_ratio = request.form.get("snap_ratio", "16:9").strip() or "16:9"
        watermark_position = request.form.get("watermark_position", "bottom-right").strip() or "bottom-right"
        watermark_opacity = request.form.get("watermark_opacity", "45").strip() or "45"
        logo_scale = request.form.get("logo_scale", "20").strip() or "20"
        output_filename = request.form.get("output_filename", "watermarked-image").strip() or "watermarked-image"
        output_format = request.form.get("output_format", "webp").strip().lower() or "webp"
        saved_source_image = request.form.get("saved_source_image", "").strip()
        crop_x = request.form.get("crop_x", "0").strip() or "0"
        crop_y = request.form.get("crop_y", "0").strip() or "0"
        crop_width = request.form.get("crop_width", "").strip()
        crop_height = request.form.get("crop_height", "").strip()
        crop_scale = request.form.get("crop_scale", "70").strip() or "70"
        watermark_x = request.form.get("watermark_x", "85").strip() or "85"
        watermark_y = request.form.get("watermark_y", "85").strip() or "85"
        watermark_rotation = request.form.get("watermark_rotation", "0").strip() or "0"
        use_watermark = request.form.get("use_watermark") == "1"

        brand_record = get_brand_record(brand)
        if brand_record and brand_record.get("logo_path"):
            brand_logo_url = image_url(brand_record.get("logo_path", ""))

        uploaded_image = request.files.get("image_file")
        source_filename = saved_source_image
        if uploaded_image and uploaded_image.filename:
            source_filename = ""

        if source_filename:
            source_image_name = Path(source_filename).name
            source_image_url = image_url(f"image_tools/{source_filename}")

        if (not uploaded_image or not uploaded_image.filename) and not source_filename:
            error = "Please upload the image you want to process."
        elif output_format not in {"png", "jpg", "jpeg", "webp"}:
            error = "Please choose PNG, JPG, JPEG, or WEBP as the export format."
        elif use_watermark and not brand:
            error = "Please select or enter a brand to use a watermark."
        elif use_watermark and not brand_record:
            error = "That brand is not saved yet. Add it first on the Brands page."
        elif use_watermark and not brand_record.get("logo_path"):
            error = "This brand does not have a logo yet. Upload one on the Brands page first."
        else:
            try:
                from PIL import Image

                if uploaded_image and uploaded_image.filename:
                    source_filename = save_uploaded_image(uploaded_image, IMAGE_TOOL_DIR, "source")
                if not source_filename:
                    raise ValueError("Please upload the image you want to process.")
                source_path = IMAGE_TOOL_DIR / source_filename
                if not source_path.exists():
                    raise ValueError("The last uploaded image could not be found. Please upload it again.")

                source_image_name = Path(source_filename).name
                source_image_url = image_url(f"image_tools/{source_filename}")
                clean_base_name = secure_filename(Path(output_filename).stem).replace("_", " ") or "watermarked-image"

                normalized_format = "jpg" if output_format == "jpeg" else output_format

                with Image.open(source_path) as source_image:
                    working_image = source_image.convert("RGBA")
                    working_image = crop_image_to_box(
                        working_image,
                        crop_x,
                        crop_y,
                        crop_width,
                        crop_height,
                    )
                    output_width, output_height = calculate_output_dimensions(
                        pixel_width,
                        pixel_height,
                        snap_ratio,
                        working_image.width,
                        working_image.height,
                    )
                    if (working_image.width, working_image.height) != (output_width, output_height):
                        working_image = working_image.resize(
                            (output_width, output_height),
                            resample=Image.Resampling.LANCZOS,
                        )
                    if use_watermark:
                        logo_path = UPLOAD_ROOT / brand_record["logo_path"]
                        with Image.open(logo_path) as logo_image:
                            working_image = apply_logo_watermark(
                                working_image,
                                logo_image,
                                watermark_position,
                                watermark_opacity,
                                logo_scale,
                                watermark_x,
                                watermark_y,
                                watermark_rotation,
                            )

                    if normalized_format in {"jpg", "webp"}:
                        working_image = working_image.convert("RGB")

                    output_name = f"{clean_base_name}.{normalized_format}"
                    output_path = IMAGE_TOOL_DIR / output_name
                    save_format = "JPEG" if normalized_format == "jpg" else normalized_format.upper()
                    save_kwargs = {"format": save_format}
                    if save_format == "JPEG":
                        save_kwargs["quality"] = 92
                    if save_format == "WEBP":
                        save_kwargs["quality"] = 92
                    working_image.save(output_path, **save_kwargs)

                result_image_url = image_url(f"image_tools/{output_name}")
                result_download_name = f"{clean_base_name}.{normalized_format}"
                success = f"Image processed as {result_download_name}."
            except ImportError:
                error = "Image processing needs Pillow. Install it with: pip install pillow"
            except ValueError as exc:
                error = str(exc)
            except Exception:
                logger.exception("image_tools action failed")
                error = "An error occurred while processing the image. Check logs/app.log for details."

    return render_template(
        "image_tools.html",
        provider=PROVIDER,
        model=MODEL,
        brand_names=brand_names,
        brand=brand,
        brand_logo_url=brand_logo_url,
        source_image_url=source_image_url,
        source_image_name=source_image_name,
        result_image_url=result_image_url,
        result_download_name=result_download_name,
        error=error,
        success=success,
        pixel_width=pixel_width,
        pixel_height=pixel_height,
        snap_ratio=snap_ratio,
        crop_x=crop_x,
        crop_y=crop_y,
        crop_width=crop_width,
        crop_height=crop_height,
        crop_scale=crop_scale,
        watermark_x=watermark_x,
        watermark_y=watermark_y,
        watermark_rotation=watermark_rotation,
        use_watermark=use_watermark,
        watermark_position=watermark_position,
        watermark_opacity=watermark_opacity,
        logo_scale=logo_scale,
        output_filename=output_filename,
        output_format=output_format,
    )


@app.route("/brands", methods=["GET", "POST"])
def brands():
    brand_name = ""
    website = ""
    niche = ""
    main_keywords = ""
    tone = ""
    notes = ""
    logo_path = ""
    check_brand = ""
    check_keyword = ""
    keyword_check_result = None
    error = None
    success = None

    edit_brand = request.args.get("edit", "").strip()
    if request.method == "GET" and edit_brand:
        brand_record = get_brand_record(edit_brand)
        if brand_record:
            brand_name = brand_record.get("name", "")
            website = brand_record.get("website", "")
            niche = brand_record.get("niche", "")
            main_keywords = brand_record.get("main_keywords", "")
            tone = brand_record.get("tone", "")
            notes = brand_record.get("notes", "")
            logo_path = brand_record.get("logo_path", "")

    if request.method == "POST":
        action = request.form.get("action", "save_brand").strip()

        if action == "save_brand":
            brand_name = request.form.get("brand_name", "").strip()
            website = request.form.get("website", "").strip()
            niche = request.form.get("niche", "").strip()
            main_keywords = request.form.get("main_keywords", "").strip()
            tone = request.form.get("tone", "").strip()
            notes = request.form.get("notes", "").strip()
            logo_upload = request.files.get("logo_file")

            if not brand_name:
                error = "Please enter a brand name."
            else:
                try:
                    if logo_upload and logo_upload.filename:
                        logo_path = f"brand_logos/{save_uploaded_image(logo_upload, BRAND_LOGO_DIR, 'logo')}"

                    upsert_brand(
                        brand_name,
                        website=website,
                        tone=tone,
                        notes=notes,
                        niche=niche,
                        main_keywords=main_keywords,
                        logo_path=logo_path,
                    )
                    success = f"Saved brand: {brand_name}"
                    brand_name = ""
                    website = ""
                    niche = ""
                    main_keywords = ""
                    tone = ""
                    notes = ""
                    logo_path = ""
                except ValueError as exc:
                    error = str(exc)
                except Exception:
                    logger.exception("brands save action failed")
                    error = "An error occurred while saving the brand. Check logs/app.log for details."

        elif action == "check_keyword":
            check_brand = request.form.get("check_brand", "").strip()
            check_keyword = request.form.get("check_keyword", "").strip()

            if not check_brand or not check_keyword:
                error = "Please enter both a brand and a keyword to check."
            else:
                try:
                    keyword_check_result = check_keyword_usage(check_brand, check_keyword)
                except Exception:
                    logger.exception("brands check_keyword action failed")
                    error = "An error occurred while checking the keyword. Check logs/app.log for details."

    return render_template(
        "brands.html",
        provider=PROVIDER,
        model=MODEL,
        brand_name=brand_name,
        website=website,
        niche=niche,
        main_keywords=main_keywords,
        tone=tone,
        notes=notes,
        logo_path=logo_path,
        logo_url=image_url(logo_path),
        check_brand=check_brand,
        check_keyword=check_keyword,
        keyword_check_result=keyword_check_result,
        brands=list_brand_records(),
        error=error,
        success=success,
    )


@app.route("/preview", methods=["POST"])
def preview():
    title = request.form.get("selected_title", "")
    keyword = request.form.get("keyword", "")
    supporting_keyword = request.form.get("supporting_keyword", "")
    meta_description = request.form.get("meta_description", "")
    content_html = request.form.get("content_html", "")

    return render_template(
        "preview.html",
        title=title,
        keyword=keyword,
        supporting_keyword=supporting_keyword,
        meta_description=meta_description,
        content_html=content_html,
    )


@app.route("/download_doc", methods=["POST"])
def download_doc():
    title = request.form.get("selected_title", "")
    keyword = request.form.get("keyword", "")
    supporting_keyword = request.form.get("supporting_keyword", "")
    meta_description = request.form.get("meta_description", "")
    content_html = request.form.get("content_html", "")

    # Create a new Document
    doc = Document()
    
    # Add title
    title_para = doc.add_paragraph(title, style='Heading 1')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Keyword: {keyword}")
    if supporting_keyword:
        doc.add_paragraph(f"Supporting Keyword: {supporting_keyword}")
    doc.add_paragraph(f"Meta Description: {meta_description}")
    
    # Add separator
    doc.add_paragraph()
    
    # Add content with formatting
    html_to_docx_paragraph(doc, content_html)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    filename = title.replace(" ", "_")[:50] or "blog_post"
    
    response = make_response(doc_io.getvalue())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
    return response


if __name__ == "__main__":
    host = os.getenv("APP_HOST", "127.0.0.1").strip() or "127.0.0.1"
    port = int(os.getenv("APP_PORT", "3444"))
    debug = os.getenv("FLASK_DEBUG", "true").strip().lower() in {"1", "true", "yes", "on"}
    app.run(host=host, port=port, debug=debug)
